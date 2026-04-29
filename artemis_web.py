#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════╗
║   ARTEMIS WEB — External Pentest Automation (Web Mode)   ║
╚══════════════════════════════════════════════════════════╝
"""

import os
import sys
import subprocess
import shutil
import threading
import re
import queue
import json
import secrets
import string
import urllib.request
import urllib.error
import urllib.parse
from datetime import datetime, timedelta
from pathlib import Path
from functools import wraps
from flask import (Flask, render_template, request, jsonify,
                   Response, send_file, session, redirect, url_for, flash)

# ── Rate limiting & CSRF store ─────────────────────────────────────────────
import time as _time
_login_attempts: dict = {}   # {ip: {"count": int, "lockout_until": float}}
_csrf_tokens:    dict = {}   # {session_id: token}  (lightweight CSRF)
LOGIN_MAX_ATTEMPTS = 5
LOGIN_LOCKOUT_SECS = 300     # 5 minutes

# ── Root check ────────────────────────────────────────────────────────────
if os.geteuid() != 0:
    print("\n[!] Artemis requires root privileges.")
    print("    Run: sudo python3 artemis_web.py\n")
    sys.exit(1)

# ── bcrypt — install if missing ───────────────────────────────────────────
try:
    import bcrypt
except ImportError:
    subprocess.run(["pip3", "install", "bcrypt", "--break-system-packages"],
                   capture_output=True)
    import bcrypt

app = Flask(__name__)

# ══════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════
RESULTS_BASE    = Path("/opt/artemis/results")
QUICK_SCANS_BASE = RESULTS_BASE / "Quick_Scans"
HISTORY_FILE   = Path("/opt/artemis/history.json")
USERS_FILE     = Path("/opt/artemis/users.json")
REQUESTS_FILE  = Path("/opt/artemis/requests.json")
WIKI_FILE      = Path("/opt/artemis/wiki.json")
CLIENTS_FILE   = Path("/opt/artemis/clients.json")
FINDINGS_FILE  = Path("/opt/artemis/findings.json")
AUDIT_LOG_FILE = Path("/opt/artemis/audit.log")
RESULTS_BASE.mkdir(parents=True, exist_ok=True)
QUICK_SCANS_BASE.mkdir(parents=True, exist_ok=True)

app.secret_key        = os.environ.get("ARTEMIS_SECRET", "artemis-secret-key-change-me")
SESSION_LIFETIME_HOURS = 3
OATHNET_API_KEY       = os.environ.get("OATHNET_API_KEY", "")

# ══════════════════════════════════════════════════════════════════════════
# TOOL CONSTANTS
# ══════════════════════════════════════════════════════════════════════════
APT_PACKAGES = {
    "nmap":         "nmap",
    "nikto":        "nikto",
    "sslscan":      "sslscan",
    "dnsenum":      "dnsenum",
    "curl":         "curl",
    "wpscan":       "wpscan",
}

GO_TOOLS = {
    "nuclei":      "github.com/projectdiscovery/nuclei/v3/cmd/nuclei@latest",
    "ffuf":        "github.com/ffuf/ffuf/v2@latest",
    "assetfinder": "github.com/tomnomnom/assetfinder@latest",
    "gowitness":   "github.com/sensepost/gowitness@latest",
    "termshot":    "github.com/homeport/termshot/cmd/termshot@latest",
    "amass":       "github.com/owasp-amass/amass/v4/...@master",
}

PIP_TOOLS = {
    "metagoofil": "metagoofil",
    "bbot":       "bbot",
}

SHCHECK_SEARCH_PATHS = [
    "/opt/shcheck/shcheck.py", "/usr/local/bin/shcheck.py",
    "/root/shcheck/shcheck.py",
    str(Path.home() / "shcheck" / "shcheck.py"),
    str(Path.home() / "tools"   / "shcheck.py"),
]
SPOOFY_SEARCH_PATHS = [
    "/opt/spoofy/spoofy.py", "/root/spoofy/spoofy.py",
    str(Path.home() / "tools"  / "spoofy.py"),
    str(Path.home() / "spoofy" / "spoofy.py"),
]
O365SCAN_SEARCH_PATHS = [
    "/opt/o365spray/o365spray.py", "/root/o365spray/o365spray.py",
    str(Path.home() / "tools"     / "o365spray.py"),
    str(Path.home() / "o365spray" / "o365spray.py"),
]

# ══════════════════════════════════════════════════════════════════════════
# GLOBAL SCAN STATE
# ══════════════════════════════════════════════════════════════════════════
scan_state = {
    "running":          False,
    "cancel_requested": False,
    "log_queue":        queue.Queue(),
    "results":          {},
    "client_folder":    None,
    "tool_paths":       {},
    "started":          "",
    "completed":        "",
}

# ══════════════════════════════════════════════════════════════════════════
# HISTORY PERSISTENCE  (last 5 assessments per user)
# ══════════════════════════════════════════════════════════════════════════

def audit_log(event: str, username: str = "", detail: str = ""):
    """Append a line to the audit log."""
    ts   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ip   = request.remote_addr if request else "—"
    line = f"[{ts}] {event:<20} user={username:<20} ip={ip:<16} {detail}\n"
    try:
        with open(AUDIT_LOG_FILE, "a") as f:
            f.write(line)
    except Exception:
        pass


def load_history() -> list:
    if HISTORY_FILE.exists():
        try:
            return json.loads(HISTORY_FILE.read_text())
        except Exception:
            pass
    return []


def load_history_for_user(username: str) -> list:
    """Return only the assessments belonging to this user."""
    return [h for h in load_history() if h.get("username") == username]


def save_assessment(client: str, date: str, domain: str,
                    folder: str, phases: dict,
                    username: str = "",
                    started: str = "",
                    completed: str = "",
                    client_id: str = "",
                    findings_count: int = 0):
    """Append assessment to history, keep only the last 5 per user."""
    history = load_history()
    entry = {
        "client":         client,
        "date":           date,
        "domain":         domain or "—",
        "folder":         folder,
        "phases":         phases,
        "username":       username,
        "started":        started,
        "completed":      completed or datetime.now().strftime("%Y-%m-%d %H:%M"),
        "client_id":      client_id,
        "findings_count": findings_count,
    }
    # Remove any existing entry for this folder, prepend new, keep 5 per user
    history = [h for h in history if h.get("folder") != folder]
    history = [entry] + history
    HISTORY_FILE.write_text(json.dumps(history, indent=2))

    # Link assessment to client record
    if client_id:
        clients = load_clients()
        if client_id in clients:
            existing = clients[client_id].get("assessments", [])
            if folder not in existing:
                existing.append(folder)
            clients[client_id]["assessments"] = existing
            clients[client_id]["last_tested"] = date
            save_clients(clients)


# ══════════════════════════════════════════════════════════════════════════
# AUTH HELPERS
# ══════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════
# TOOL HELPERS
# ══════════════════════════════════════════════════════════════════════════

def _add_to_path(directory: str, log_fn):
    if directory and directory not in os.environ.get("PATH", "").split(":"):
        os.environ["PATH"] = os.environ.get("PATH", "") + f":{directory}"
        log_fn(f"↳ Added to PATH: {directory}", "dim")


def _find_tool_on_disk(tool: str) -> str | None:
    SEARCH_DIRS = [
        "/root/go/bin", "/usr/local/bin", "/usr/bin",
        "/opt/metasploit-framework/bin", "/opt/metasploit/bin",
    ]
    for d in SEARCH_DIRS:
        candidate = Path(d) / tool
        if candidate.exists() and os.access(str(candidate), os.X_OK):
            return d
    try:
        result = subprocess.run(
            ["find", "/usr", "/opt", "/root", "/home",
             "-name", tool, "-type", "f", "-maxdepth", "8"],
            capture_output=True, text=True, timeout=15)
        for line in result.stdout.splitlines():
            line = line.strip()
            if line and os.access(line, os.X_OK):
                return str(Path(line).parent)
    except Exception:
        pass
    return None


def _find_script(search_paths: list, name: str) -> str | None:
    for p in search_paths:
        if Path(p).exists():
            return p
    try:
        result = subprocess.run(
            ["find", "/", "-name", name, "-maxdepth", "7"],
            capture_output=True, text=True, timeout=10)
        lines = [l.strip() for l in result.stdout.splitlines() if l.strip()]
        return lines[0] if lines else None
    except Exception:
        return None


def _tor_running() -> bool:
    """Check if Tor is listening on port 9050."""
    import socket
    try:
        with socket.create_connection(("127.0.0.1", 9050), timeout=2):
            return True
    except Exception:
        return False


def find_shcheck() -> str | None:
    for p in SHCHECK_SEARCH_PATHS:
        if Path(p).exists():
            return p
    try:
        result = subprocess.run(
            ["find", "/", "-name", "shcheck.py", "-maxdepth", "6"],
            capture_output=True, text=True, timeout=10)
        lines = [l.strip() for l in result.stdout.splitlines() if l.strip()]
        return lines[0] if lines else None
    except Exception:
        return None


def check_and_install_tools(log_fn):
    log_fn("🔍 Checking and installing required tools...", "info")

    # ── 1. Ensure Go bin dirs are in PATH first ───────────────────────
    for go_bin in ["/root/go/bin", "/usr/local/go/bin",
                   str(Path.home() / "go" / "bin")]:
        if Path(go_bin).is_dir():
            _add_to_path(go_bin, log_fn)

    # ── 2. Ensure pip/local bin dirs are in PATH ──────────────────────
    for pip_bin in ["/usr/local/bin", "/usr/bin",
                    str(Path.home() / ".local" / "bin")]:
        if Path(pip_bin).is_dir():
            _add_to_path(pip_bin, log_fn)

    # ── 3. APT packages ───────────────────────────────────────────────
    log_fn("  [apt] Checking system packages...", "info")
    apt_needed = [pkg for tool, pkg in APT_PACKAGES.items()
                  if shutil.which(tool) is None]
    if apt_needed:
        log_fn(f"  [apt] Installing: {', '.join(apt_needed)}", "warn")
        subprocess.run(["apt-get", "update", "-qq"], capture_output=True)
        subprocess.run(["apt-get", "install", "-y"] + apt_needed,
                       capture_output=True)
    for tool, pkg in APT_PACKAGES.items():
        found = shutil.which(tool) or _find_tool_on_disk(tool)
        if found:
            _add_to_path(str(Path(found).parent), log_fn)
            log_fn(f"  ✓ {tool}", "success")
        else:
            log_fn(f"  ✗ {tool} — apt install failed, check manually", "error")

    # ── 4. Go tools ───────────────────────────────────────────────────
    log_fn("  [go] Checking Go tools...", "info")
    go_path = shutil.which("go")
    for tool, module in GO_TOOLS.items():
        found = shutil.which(tool) or _find_tool_on_disk(tool)
        if not found:
            if go_path:
                log_fn(f"  [go] Installing {tool}...", "warn")
                r = subprocess.run(
                    ["go", "install", module],
                    capture_output=True, text=True,
                    env={**os.environ,
                         "GOPATH": "/root/go",
                         "PATH":   os.environ.get("PATH", "") + ":/root/go/bin"})
                if r.returncode != 0:
                    log_fn(f"  ✗ {tool} — go install failed: {r.stderr[:100]}", "error")
                    continue
            else:
                log_fn(f"  ✗ {tool} — Go not installed, skipping", "error")
                continue
        found = shutil.which(tool) or _find_tool_on_disk(tool)
        if found:
            _add_to_path(str(Path(found).parent), log_fn)
            log_fn(f"  ✓ {tool}", "success")
        else:
            log_fn(f"  ✗ {tool} — not found after go install", "error")

    # ── 5. Pip tools ──────────────────────────────────────────────────
    log_fn("  [pip] Checking Python tools...", "info")

    # metagoofil and bbot — search multiple locations
    pip_tool_map = {
        "metagoofil": {"pkg": "metagoofil", "bins": ["metagoofil"]},
        "bbot":       {"pkg": "bbot",       "bins": ["bbot"]},
    }

    for tool, info in pip_tool_map.items():
        # Check all possible binary names
        found = None
        for bin_name in info["bins"]:
            found = shutil.which(bin_name)
            if found:
                break
            # Also check common pip install locations explicitly
            for pip_dir in ["/usr/local/bin", "/usr/bin",
                            str(Path.home() / ".local" / "bin"),
                            "/root/.local/bin"]:
                candidate = Path(pip_dir) / bin_name
                if candidate.exists():
                    found = str(candidate)
                    _add_to_path(pip_dir, log_fn)
                    break
            if found:
                break

        if not found:
            log_fn(f"  [pip] Installing {tool} ({info['pkg']})...", "warn")
            subprocess.run(
                ["pip3", "install", info["pkg"], "--break-system-packages",
                 "--quiet"],
                capture_output=True)
            # Re-check after install
            for bin_name in info["bins"]:
                for pip_dir in ["/usr/local/bin", "/usr/bin",
                                str(Path.home() / ".local" / "bin"),
                                "/root/.local/bin"]:
                    candidate = Path(pip_dir) / bin_name
                    if candidate.exists():
                        found = str(candidate)
                        _add_to_path(pip_dir, log_fn)
                        break
                if found:
                    break

        if found:
            log_fn(f"  ✓ {tool} ({found})", "success")
        else:
            log_fn(f"  ✗ {tool} — pip install failed or binary not found", "error")

    # ── 6. Git-cloned tools — auto-clone if missing ───────────────────
    log_fn("  [git] Checking git tools...", "info")

    git_tools = {
        "shcheck": {
            "paths":  SHCHECK_SEARCH_PATHS,
            "repo":   "https://github.com/santoru/shcheck",
            "dest":   "/opt/shcheck",
            "script": "shcheck.py",
            "req":    "/opt/shcheck/requirements.txt",
        },
        "spoofy": {
            "paths":  SPOOFY_SEARCH_PATHS,
            "repo":   "https://github.com/MattKeeley/Spoofy",
            "dest":   "/opt/spoofy",
            "script": "spoofy.py",
            "req":    "/opt/spoofy/requirements.txt",
        },
        "o365spray": {
            "paths":  O365SCAN_SEARCH_PATHS,
            "repo":   "https://github.com/0xZDH/o365spray",
            "dest":   "/opt/o365spray",
            "script": "o365spray.py",
            "req":    "/opt/o365spray/requirements.txt",
        },
    }

    results = {}
    for name, info in git_tools.items():
        found = _find_script(info["paths"], info["script"])
        if not found:
            dest = Path(info["dest"])
            if not dest.exists():
                log_fn(f"  [git] Cloning {name}...", "warn")
                r = subprocess.run(
                    ["git", "clone", info["repo"], str(dest)],
                    capture_output=True, text=True)
                if r.returncode != 0:
                    log_fn(f"  ✗ {name} — git clone failed: {r.stderr[:100]}", "error")
                    results[name] = None
                    continue
                # Install requirements if present
                req = Path(info["req"])
                if req.exists():
                    subprocess.run(
                        ["pip3", "install", "-r", str(req),
                         "--break-system-packages", "--quiet"],
                        capture_output=True)
            found = _find_script(info["paths"], info["script"])

        if found:
            log_fn(f"  ✓ {name} ({found})", "success")
        else:
            log_fn(f"  ✗ {name} — not found after clone attempt", "error")
        results[name] = found

    # ── 7. Metasploit — manual install only, just check ───────────────
    msf = shutil.which("msfconsole") or _find_tool_on_disk("msfconsole")
    if msf:
        _add_to_path(str(Path(msf).parent), log_fn)
        log_fn("  ✓ msfconsole", "success")
    else:
        log_fn("  ✗ msfconsole — install manually from rapid7", "warn")

    # ── 8. Update nuclei templates ────────────────────────────────────
    if shutil.which("nuclei"):
        log_fn("  Updating nuclei templates...", "info")
        subprocess.run(["nuclei", "-update-templates"],
                       capture_output=True, timeout=120)
        log_fn("  ✓ nuclei templates updated", "success")

    log_fn("✅ Tool check complete.", "success")

    # ── 9. Functional tests ───────────────────────────────────────────
    log_fn("🧪 Running functional tests...", "info")
    _run_functional_tests(log_fn)

    return {
        "shcheck":  results.get("shcheck"),
        "spoofy":   results.get("spoofy"),
        "o365scan": results.get("o365spray"),
    }


def _run_functional_tests(log_fn):
    """
    Execute each tool with a harmless version/help flag to verify it actually
    runs — not just that the binary exists in PATH. Catches broken installs
    like Nikto's missing nikto.pl early, before a real scan starts.
    """

    def test(label, cmd, success_hint=None):
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
            output = (r.stdout + r.stderr).lower()
            # Most version/help commands exit 0 or 1; 255 is also common for --help
            if r.returncode not in (0, 1, 255):
                snippet = (r.stderr or r.stdout)[:120].strip()
                log_fn(f"  ✗ {label} — exited {r.returncode}: {snippet}", "error")
                return False
            if success_hint and success_hint.lower() not in output:
                snippet = (r.stdout + r.stderr)[:120].strip()
                log_fn(f"  ✗ {label} — unexpected output: {snippet}", "error")
                return False
            log_fn(f"  ✓ {label} functional", "success")
            return True
        except FileNotFoundError:
            log_fn(f"  ✗ {label} — binary not found in PATH", "error")
            return False
        except subprocess.TimeoutExpired:
            log_fn(f"  ✗ {label} — timed out on version check", "error")
            return False
        except Exception as ex:
            log_fn(f"  ✗ {label} — {ex}", "error")
            return False

    test("nmap",         ["nmap",         "--version"],  "nmap")
    test("nikto",        ["nikto",        "-Version"],   "nikto")
    test("sslscan",      ["sslscan",      "--version"],  "sslscan")
    test("dnsenum",      ["dnsenum",      "--help"],     "dnsenum")
    test("curl",         ["curl",         "--version"],  "curl")
    test("nuclei",       ["nuclei",       "-version"],   "nuclei")
    test("ffuf",         ["ffuf",         "-V"],         "ffuf")
    test("assetfinder",  ["assetfinder",  "--help"],     "assetfinder")
    test("gowitness",    ["gowitness",    "--help"],     "gowitness")
    test("termshot",     ["termshot",     "--help"],     "termshot")
    test("msfconsole",   ["msfconsole",  "--version"],  "metasploit")
    test("bbot",         ["bbot",         "--help"],     "bbot")
    test("wpscan",       ["wpscan",       "--version"],  "wpscan")
    test("amass",        ["amass",        "-version"],   "amass")
    test("masscan",      ["masscan",      "--version"],  "masscan")

    log_fn("🧪 Functional tests complete.", "info")

def _write_msf_rc(path: Path, scope_list: list[str]):
    rhosts = " ".join(scope_list)
    path.write_text(f"""# Auto-generated by Artemis
setg RHOSTS {rhosts}
setg THREADS 10
use auxiliary/scanner/http/http_version
run
use auxiliary/scanner/ssh/ssh_version
run
use auxiliary/scanner/ftp/ftp_version
run
use auxiliary/scanner/smtp/smtp_version
run
use auxiliary/scanner/ssl/openssl_heartbleed
run
use auxiliary/scanner/smb/smb_version
run
use auxiliary/scanner/smb/smb_ms17_010
run
exit
""")


def run_scan(scope_list, url_list, domain, phases, tools, folder, tool_paths):
    lq = scan_state["log_queue"]

    def log(msg, tag="info"):
        ts = datetime.now().strftime("%H:%M:%S")
        lq.put({"ts": ts, "msg": msg, "tag": tag})
        if folder and folder.exists():
            with open(folder / "artemis.log", "a") as f:
                f.write(f"[{ts}] {msg}\n")

    def phase(title):
        log("─" * 52, "phase")
        log(f"  {title}", "phase")
        log("─" * 52, "phase")

    # Strip ANSI escape codes from output files
    _ansi_re = re.compile(r'\x1b\[[0-9;]*[mGKHF]|\x1b\[[0-9;]*[A-Za-z]|\x1b\(B|\x1b=|\x1b>')

    def run_tool(cmd, out_file, label, screenshot_name=None, timeout=600):
        # Check cancel before starting each tool
        if scan_state.get("cancel_requested"):
            log(f"⊘ {label} — cancelled", "warn")
            return -2
        log(f"⟶ {label}", "info")
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
            # Strip ANSI codes before writing to file
            clean_out = _ansi_re.sub('', r.stdout + r.stderr)
            with open(out_file, "w") as fh:
                fh.write(clean_out)
            if r.returncode not in (0, 1):
                log(f"⚠ {label} exited with code {r.returncode}", "warn")
            else:
                log(f"✓ {label} complete", "success")
            if screenshot_name and shutil.which("termshot"):
                try:
                    png = evidence_dir / f"{screenshot_name}.png"
                    subprocess.run(["termshot", "-f", str(png), "--"] + cmd,
                                   capture_output=True, timeout=120)
                except Exception:
                    pass
            return r.returncode
        except subprocess.TimeoutExpired:
            log(f"✗ {label} timed out after {timeout}s", "error")
            return -1
        except FileNotFoundError:
            log(f"✗ {label} — tool not found in PATH", "error")
            return 127

    shcheck_path  = tool_paths.get("shcheck")
    spoofy_path   = tool_paths.get("spoofy")
    o365scan_path = tool_paths.get("o365scan")

    evidence_dir = folder / "Assessment_Evidence"
    evidence_dir.mkdir(exist_ok=True)

    results = {
        "client":           scan_state["results"].get("client", ""),
        "date":             scan_state["results"].get("date", ""),
        "domain":           domain,
        "targets":          scope_list,
        "scope_list":       scope_list,
        "url_list":         url_list,
        "subdomains":       [],
        "open_ports":       {},
        "vulnerabilities":  {},
        "missing_headers":  {},
        "ffuf_findings":    {},
        "msf_findings":     [],
        "o365_findings":    {},
        "metagoofil":       [],
    }

    if phases.get("recon"):
        phase("PHASE 1 — Reconnaissance / OSINT")
        p1 = folder / "1_recon"; p1.mkdir(exist_ok=True)
        subdomains_found = []

        osint_targets = [domain] if domain else scope_list
        for target in osint_targets:
            safe_t = re.sub(r"[^\w\-]", "_", target)
            log(f"  OSINT target: {target}", "dim")

            if tools.get("assetfinder", True):
                run_tool(["assetfinder", "--subs-only", target],
                         p1 / f"assetfinder_{target}.txt",
                         f"assetfinder [{target}]",
                         screenshot_name=f"phase1_assetfinder_{safe_t}")
                sf = p1 / f"assetfinder_{target}.txt"
                if sf.exists():
                    subdomains_found.extend(
                        [l.strip() for l in sf.read_text().splitlines() if l.strip()])
            else:
                log("  — assetfinder skipped", "dim")

            if tools.get("dnsenum", True):
                run_tool(["dnsenum", "--enum", "--noreverse", target],
                         p1 / f"dnsenum_{target}.txt",
                         f"dnsenum [{target}]",
                         screenshot_name=f"phase1_dnsenum_{safe_t}")
            else:
                log("  — dnsenum skipped", "dim")

        if tools.get("curl_sweep", True):
            log("⟶ curl security header sweep...", "info")
            sweep_file = p1 / "security_headers_sweep.txt"
            with open(sweep_file, "w") as hsf:
                for url in url_list:
                    hsf.write(f"\n{'='*60}\n{url}\n{'='*60}\n")
                    try:
                        r = subprocess.run(
                            ["curl", "-s", "-I", "--max-time", "10", "-L",
                             "-A", "Mozilla/5.0", url],
                            capture_output=True, text=True, timeout=15)
                        hsf.write(r.stdout)
                    except Exception as ex:
                        hsf.write(f"ERROR: {ex}\n")
            log("✓ curl header sweep complete", "success")
        else:
            log("  — curl header sweep skipped", "dim")

        results["subdomains"] = list(set(subdomains_found))

        if domain and spoofy_path and tools.get("spoofy", True):
            safe_d = re.sub(r"[^\w\-]", "_", domain)
            run_tool(["python3", spoofy_path, "-d", domain],
                     p1 / f"spoofy_{safe_d}.txt",
                     f"spoofy [{domain}]",
                     screenshot_name=f"phase1_spoofy_{safe_d}")
        elif domain and not tools.get("spoofy", True):
            log("  — spoofy skipped", "dim")
        elif domain:
            log("⚠ spoofy.py not found — skipping", "warn")

        if domain and o365scan_path and tools.get("o365spray", True):
            safe_d = re.sub(r"[^\w\-]", "_", domain)
            o365_out = p1 / f"o365scan_{safe_d}.txt"
            run_tool(["python3", o365scan_path,
                      "--validate", "--domain", domain, "--output", str(p1)],
                     o365_out, f"o365spray [{domain}]",
                     screenshot_name=f"phase1_o365scan_{safe_d}")
            if o365_out.exists():
                raw = o365_out.read_text()
                results["o365_findings"] = {
                    "domain": domain, "raw": raw[:4000],
                    "o365":     "Microsoft 365" in raw or "True" in raw,
                    "adfs":     "ADFS" in raw,
                    "exchange": "Exchange" in raw,
                }
        elif domain and not tools.get("o365spray", True):
            log("  — o365spray skipped", "dim")
        elif domain:
            log("⚠ o365spray not found — skipping", "warn")

        if domain and tools.get("metagoofil", True):
            if shutil.which("metagoofil"):
                safe_d       = re.sub(r"[^\w\-]", "_", domain)
                meta_exts    = "pdf,doc,docx,xls,xlsx,ppt,pptx"
                meta_dir     = p1 / f"metagoofil_{safe_d}"
                meta_dir.mkdir(exist_ok=True)
                meta_out     = p1 / f"metagoofil_{safe_d}.txt"

                # Use proxychains4 if available and Tor is listening on 9050
                use_proxy = (shutil.which("proxychains4") and
                             _tor_running())
                proxy_prefix = ["proxychains4", "-q"] if use_proxy else []
                if use_proxy:
                    log(f"⟶ metagoofil — root domain [{domain}] (via proxychains/Tor)", "info")
                else:
                    log(f"⟶ metagoofil — root domain [{domain}] (no proxy — install tor + proxychains4 to avoid blocks)", "warn")

                run_tool(
                    proxy_prefix + ["metagoofil", "-d", domain, "-t", meta_exts,
                     "-n", "10", "-o", str(meta_dir)],
                    meta_out,
                    f"metagoofil [{domain}]",
                    screenshot_name=f"phase1_metagoofil_{safe_d}",
                    timeout=900
                )
                meta_findings = []
                for mf in list(meta_dir.glob("*")) + [meta_out]:
                    if mf.exists() and mf.is_file():
                        meta_findings.extend([
                            l.strip() for l in mf.read_text(errors="ignore").splitlines()
                            if l.strip() and not l.startswith("#")
                        ])
                results["metagoofil"] = list(set(meta_findings))
            else:
                log("⚠ metagoofil not found — install: pip3 install metagoofil --break-system-packages", "warn")
        elif domain and not tools.get("metagoofil", True):
            log("  — metagoofil skipped", "dim")
        else:
            log("⚠ No domain — skipping metagoofil", "warn")

        # ── BBOT — passive or active depending on preset ──────────────
        if domain and tools.get("bbot", True):
            if shutil.which("bbot"):
                safe_d    = re.sub(r"[^\w\-]", "_", domain)
                bbot_out  = p1 / f"bbot_{safe_d}"
                bbot_mode = tools.get("bbot_mode", "passive")

                if bbot_mode == "active":
                    # Everything Everywhere All At Once — full aggressive scan
                    log(f"⟶ bbot ACTIVE (everything) [{domain}]", "info")
                    bbot_cmd = [
                        "bbot", "-t", domain,
                        "-p", "everything",
                        "-o", str(bbot_out),
                        "--yes",
                        "-s",
                        "-om", "human,json",
                    ]
                else:
                    # Passive only — safe modules, no direct target interaction
                    log(f"⟶ bbot PASSIVE (safe) [{domain}]", "info")
                    bbot_cmd = [
                        "bbot", "-t", domain,
                        "-f", "safe,passive",
                        "-o", str(bbot_out),
                        "--yes",
                        "-s",
                        "-om", "human,json",
                    ]

                run_tool(bbot_cmd,
                         p1 / f"bbot_{safe_d}.txt",
                         f"bbot [{bbot_mode}] [{domain}]",
                         screenshot_name=f"phase1_bbot_{safe_d}",
                         timeout=3600)

                # Parse bbot output for subdomains, emails, IPs
                bbot_txt = p1 / f"bbot_{safe_d}.txt"
                if bbot_txt.exists():
                    content = bbot_txt.read_text()
                    bbot_subs = re.findall(
                        r"\b([a-zA-Z0-9\-]+(?:\.[a-zA-Z0-9\-]+)*\." +
                        re.escape(domain) + r")\b", content)
                    results["subdomains"] = list(set(
                        results.get("subdomains", []) + bbot_subs))
                    log(f"  ↳ bbot found {len(bbot_subs)} additional subdomains", "dim")
            else:
                log("⚠ bbot not found — install: pip3 install bbot", "warn")
        elif domain and not tools.get("bbot", True):
            log("  — bbot skipped", "dim")
        else:
            log("⚠ No domain — skipping bbot", "warn")

        # ── Amass — passive or active depending on preset ─────────────
        if domain and tools.get("amass", True):
            if shutil.which("amass"):
                safe_d     = re.sub(r"[^\w\-]", "_", domain)
                amass_out  = p1 / f"amass_{safe_d}.txt"
                amass_mode = tools.get("amass_mode", "passive")

                if amass_mode == "active":
                    log(f"⟶ amass ACTIVE (brute force + zone transfer) [{domain}]", "info")
                    wordlist = "/usr/share/seclists/Discovery/DNS/sublist3r-top1mil-20000.txt"
                    if not Path(wordlist).exists():
                        wordlist = "/usr/share/seclists/Discovery/DNS/bitquark-subdomains-top100000.txt"
                    if not Path(wordlist).exists():
                        wordlist = ""
                    amass_cmd = [
                        "amass", "enum", "-active", "-nocolor",
                        "-d", domain,
                        "-o", str(amass_out),
                        "-timeout", "30",
                    ]
                    if wordlist:
                        amass_cmd += ["-brute", "-w", wordlist]
                else:
                    log(f"⟶ amass PASSIVE (OSINT sources only) [{domain}]", "info")
                    amass_cmd = [
                        "amass", "enum", "-passive",
                        "-d", domain,
                        "-o", str(amass_out),
                        "-timeout", "20",
                    ]

                run_tool(amass_cmd,
                         amass_out,
                         f"amass [{amass_mode}] [{domain}]",
                         screenshot_name=f"phase1_amass_{safe_d}", timeout=3600)

                # Parse amass output — one subdomain per line
                if amass_out.exists():
                    amass_subs = [l.strip() for l in amass_out.read_text().splitlines()
                                  if l.strip() and "." in l]
                    results["subdomains"] = list(set(
                        results.get("subdomains", []) + amass_subs))
                    log(f"  ↳ amass found {len(amass_subs)} subdomains", "dim")
            else:
                log("⚠ amass not found — install: go install github.com/owasp-amass/amass/v4/...@master", "warn")
        elif domain and not tools.get("amass", True):
            log("  — amass skipped", "dim")
        else:
            log("⚠ No domain — skipping amass", "warn")

        # ── metagoofil subdomain sweep ────────────────────────────────
        # Runs after all OSINT tools so subdomains list is fully populated
        if domain and tools.get("metagoofil", True) and shutil.which("metagoofil"):
            all_subs = results.get("subdomains", [])
            if all_subs:
                # Priority keywords — subdomains most likely to host interesting docs
                PRIORITY_KEYWORDS = [
                    "mail", "vpn", "portal", "admin", "dev", "stage", "staging",
                    "remote", "citrix", "login", "secure", "extranet", "intranet",
                    "sharepoint", "files", "docs", "hr", "finance", "it", "helpdesk",
                ]
                # Score each subdomain — more keyword matches = higher priority
                def sub_score(s):
                    s_lower = s.lower()
                    return sum(1 for k in PRIORITY_KEYWORDS if k in s_lower)

                priority_subs = sorted(
                    [s for s in all_subs if sub_score(s) > 0],
                    key=sub_score, reverse=True
                )
                # Fill remaining slots with non-priority subs
                other_subs = [s for s in all_subs if sub_score(s) == 0]
                targets = (priority_subs + other_subs)[:20]

                log(f"⟶ metagoofil subdomain sweep — {len(targets)} subdomain(s) "
                    f"({len(priority_subs)} priority, {max(0, len(targets)-len(priority_subs))} other)", "info")

                meta_exts = "pdf,doc,docx,xls,xlsx,ppt,pptx"
                use_proxy    = (shutil.which("proxychains4") and _tor_running())
                proxy_prefix = ["proxychains4", "-q"] if use_proxy else []
                import time
                for sub in targets:
                    safe_s   = re.sub(r"[^\w\-]", "_", sub)
                    sub_dir  = p1 / f"metagoofil_{safe_s}"
                    sub_dir.mkdir(exist_ok=True)
                    sub_out  = p1 / f"metagoofil_{safe_s}.txt"
                    log(f"  ↳ metagoofil [{sub}]{'  (proxychains)' if use_proxy else ''}", "dim")
                    try:
                        with open(sub_out, "w") as mf:
                            subprocess.run(
                                proxy_prefix + ["metagoofil", "-d", sub, "-t", meta_exts,
                                 "-n", "5", "-o", str(sub_dir)],
                                stdout=mf, stderr=subprocess.STDOUT,
                                text=True, timeout=300
                            )
                    except subprocess.TimeoutExpired:
                        log(f"  ✗ metagoofil timed out for {sub} — continuing", "warn")
                    except Exception as ex:
                        log(f"  ✗ metagoofil error for {sub}: {ex}", "warn")

                    # Merge findings into results
                    sub_findings = []
                    for mf in list(sub_dir.glob("*")) + [sub_out]:
                        if mf.exists() and mf.is_file():
                            sub_findings.extend([
                                l.strip() for l in mf.read_text(errors="ignore").splitlines()
                                if l.strip() and not l.startswith("#")
                            ])
                    results["metagoofil"] = list(set(
                        results.get("metagoofil", []) + sub_findings))

                    # Delay between subdomains — longer through Tor to avoid blocks
                    time.sleep(6 if use_proxy else 3)

                log(f"  ↳ metagoofil subdomain sweep complete", "success")
            else:
                log("  — metagoofil subdomain sweep: no subdomains discovered yet", "dim")

        # ── Nuclei passive — passive-tagged templates only ─────────────
        if url_list and tools.get("nuclei_passive", True):
            if shutil.which("nuclei"):
                log("⟶ nuclei PASSIVE (passive-tagged templates) — recon phase", "info")
                for url in url_list:
                    safe_t        = re.sub(r"[^\w\-]", "_", re.sub(r"^https?://", "", url).rstrip("/"))
                    nuclei_p_out  = p1 / f"nuclei_passive_{safe_t}.txt"
                    run_tool(
                        ["nuclei", "-target", url,
                         "-passive",
                         "-o", str(nuclei_p_out), "-nc"],
                        nuclei_p_out,
                        f"nuclei passive [{url}]",
                        screenshot_name=f"phase1_nuclei_passive_{safe_t}"
                    )
                    if nuclei_p_out.exists():
                        lines = [l.strip() for l in nuclei_p_out.read_text().splitlines() if l.strip()]
                        if lines:
                            log(f"  ↳ nuclei passive found {len(lines)} result(s) for {url}", "dim")
            else:
                log("⚠ nuclei not found — skipping passive nuclei", "warn")
        elif url_list and not tools.get("nuclei_passive", True):
            log("  — nuclei passive skipped", "dim")
        else:
            log("⚠ No URLs — skipping nuclei passive", "warn")

        # ── OathNet breach credential check ───────────────────────────
        if tools.get("breach_check", True):
            api_key = OATHNET_API_KEY
            if not api_key:
                log("⚠ OATHNET_API_KEY not set — skipping breach check", "warn")
                log("  Set it in /etc/systemd/system/artemis.service and restart", "dim")
            else:
                # Build email targets:
                # 1. Query domain directly (catches all breached emails for the domain)
                # 2. Also check any specific emails discovered by bbot
                discovered_emails = set()

                # From bbot output files
                for f in p1.glob("bbot_*.txt"):
                    discovered_emails.update(re.findall(
                        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
                        f.read_text(errors="ignore")
                    ))

                # Filter to target domain emails only
                if domain:
                    domain_emails = {e for e in discovered_emails
                                     if e.lower().endswith(f"@{domain.lower()}")}
                    discovered_emails = domain_emails

                # Always include the domain itself as a search target
                # OathNet v2 API accepts email or domain as q parameter
                targets = []
                if domain:
                    targets.append(domain)          # domain search: example.com
                targets.extend(sorted(discovered_emails))

                if not targets:
                    log("  — breach check: no domain or emails to check", "dim")
                else:
                    log(f"⟶ OathNet breach check — domain: {domain}"
                        f"{f' + {len(discovered_emails)} email(s)' if discovered_emails else ''}", "info")
                    breach_out  = p1 / "breach_check.txt"
                    breach_hits = []
                    breach_errors = 0

                    with open(breach_out, "w") as bf:
                        bf.write(f"OathNet Breach Credential Check\n")
                        bf.write(f"Checked: {datetime.now().isoformat()}\n")
                        bf.write(f"Domain: {domain}\n")
                        bf.write(f"Individual emails checked: {len(discovered_emails)}\n")
                        bf.write("=" * 60 + "\n\n")

                        for target in targets:
                            try:
                                url_req = (
                                    f"https://oathnet.org/api/service/v2/breach/search"
                                    f"?q={urllib.parse.quote(target, safe='@._-')}"
                                )
                                log(f"  → querying OathNet for: {target}", "dim")
                                req = urllib.request.Request(
                                    url_req,
                                    headers={"x-api-key": api_key,
                                             "User-Agent": "Artemis-PenTest/1.0"}
                                )
                                with urllib.request.urlopen(req, timeout=15) as resp:
                                    body = json.loads(resp.read().decode())

                                results_found = body.get("data", {}).get("meta", {}).get("total", 0)
                                items         = body.get("data", {}).get("items", [])
                                sources       = list({i.get("dbname", "unknown")
                                                      for i in items if isinstance(i, dict)})

                                if results_found > 0:
                                    bf.write(f"[FOUND] {target}\n")
                                    bf.write(f"  Breach records: {results_found}\n")
                                    if sources:
                                        bf.write(f"  Sources: {', '.join(sources)}\n")
                                    bf.write("\n")
                                    breach_hits.append({
                                        "target":  target,
                                        "count":   results_found,
                                        "sources": sources,
                                    })
                                    log(f"  ⚠ BREACH: {target} — {results_found} record(s)"
                                        f" in {', '.join(sources[:3])}", "warn")
                                else:
                                    bf.write(f"[CLEAN] {target}\n")
                                    log(f"  ✓ clean: {target}", "dim")

                            except urllib.error.HTTPError as e:
                                if e.code == 401:
                                    log("✗ OathNet API key invalid or expired", "error")
                                    bf.write(f"[ERROR] {target} — API key invalid\n")
                                    breach_errors += 1
                                    break
                                elif e.code == 429:
                                    log("✗ OathNet rate limit hit — stopping breach check", "warn")
                                    bf.write(f"[RATE_LIMITED] Stopped at {target}\n")
                                    break
                                else:
                                    log(f"  ✗ HTTP {e.code} for {target}", "error")
                                    bf.write(f"[ERROR] {target} — HTTP {e.code}\n")
                                    breach_errors += 1
                            except Exception as ex:
                                log(f"  ✗ breach check error for {target}: {ex}", "error")
                                bf.write(f"[ERROR] {target} — {ex}\n")
                                breach_errors += 1

                        bf.write("\n" + "=" * 60 + "\n")
                        bf.write(f"Summary: {len(breach_hits)} breach(es) found, "
                                 f"{len(targets) - len(breach_hits) - breach_errors} clean, "
                                 f"{breach_errors} error(s)\n")

                    results["breach_hits"] = breach_hits
                    if breach_hits:
                        log(f"  ↳ {len(breach_hits)} breach(es) found — see breach_check.txt", "warn")
                    else:
                        log(f"  ↳ No breaches found for {domain}", "success")
        else:
            log("  — breach check skipped", "dim")

    if phases.get("scan"):
        phase("PHASE 2 — Port & Service Scanning")
        p2 = folder / "2_scan"; p2.mkdir(exist_ok=True)
        open_ports     = {}
        gowitness_urls = []

        for target in scope_list:
            safe_t = re.sub(r"[^\w\-]", "_", target)
            log(f"  Scanning target: {target}", "dim")

            # ── Step 1: Masscan — fast full-port discovery ─────────────
            masscan_ports = []
            if tools.get("nmap_tcp", True) and tools.get("masscan", True) and shutil.which("masscan"):
                log(f"⟶ masscan — full port sweep [{target}]", "info")
                masscan_out = p2 / f"masscan_{safe_t}.txt"
                try:
                    r = subprocess.run(
                        ["masscan", target, "-p", "0-65535",
                         "--rate", "1000",
                         "--wait", "3",
                         "-oL", str(masscan_out)],
                        capture_output=True, text=True, timeout=600
                    )
                    if masscan_out.exists():
                        # masscan -oL format: "open tcp 80 1.2.3.4 ..."
                        masscan_ports = list(set(re.findall(
                            r"^open\s+tcp\s+(\d+)", masscan_out.read_text(),
                            re.MULTILINE)))
                        masscan_ports.sort(key=int)
                        log(f"  ↳ masscan found {len(masscan_ports)} open TCP port(s)", "dim")
                except subprocess.TimeoutExpired:
                    log(f"✗ masscan timed out for {target}", "error")
                except FileNotFoundError:
                    log("⚠ masscan not found — falling back to nmap -p-", "warn")
                except Exception as ex:
                    log(f"✗ masscan error: {ex}", "error")
            elif tools.get("nmap_tcp", True) and not tools.get("masscan", True):
                log(f"  — masscan skipped — nmap will scan all ports directly", "dim")
            elif tools.get("nmap_tcp", True) and not shutil.which("masscan"):
                log("⚠ masscan not installed — nmap will scan all ports directly", "warn")

            # ── Step 2: Nmap TCP — service detection on discovered ports ─
            if tools.get("nmap_tcp", True):
                # If masscan found ports, scan only those; otherwise fall back to -p-
                if masscan_ports:
                    port_arg = ",".join(masscan_ports)
                    log(f"⟶ nmap TCP — service scan on {len(masscan_ports)} port(s) [{target}]", "info")
                    nmap_cmd = ["nmap", "-sV", "-sC", "--open",
                                "-T4", "-p", port_arg,
                                "-oN", str(p2 / f"nmap_tcp_{target}.txt"), target]
                else:
                    log(f"⟶ nmap TCP — full port scan [{target}]", "info")
                    nmap_cmd = ["nmap", "-sS", "-sV", "-sC", "-p-", "--open",
                                "-T4", "--min-rate", "1000",
                                "-oN", str(p2 / f"nmap_tcp_{target}.txt"), target]

                run_tool(nmap_cmd,
                         p2 / f"nmap_tcp_{target}.txt",
                         f"nmap TCP [{target}]",
                         screenshot_name=f"phase2_nmap_tcp_{safe_t}", timeout=3600)

                txt_file = p2 / f"nmap_tcp_{target}.txt"
                ports = []
                if txt_file.exists():
                    ports = re.findall(r"(\d+)/tcp\s+open", txt_file.read_text())
                    open_ports[target] = ports
                    web_ports = {"80","443","8080","8443","8000","8888","9090","3000","4443","4080"}
                    for p in ports:
                        scheme = "https" if p in {"443","8443","4443"} else "http"
                        if p in web_ports:
                            gowitness_urls.append(f"{scheme}://{target}:{p}")
                        elif p not in {"80","443"}:
                            gowitness_urls.append(f"http://{target}:{p}")
                    gowitness_urls += [f"http://{target}", f"https://{target}"]
            else:
                log(f"  — nmap TCP skipped for {target}", "dim")

            if tools.get("nmap_udp", True):
                run_tool(["nmap", "-sU", "--top-ports", "200", "-T4",
                          "-oN", str(p2 / f"nmap_udp_{target}.txt"), target],
                         p2 / f"nmap_udp_{target}.txt",
                         f"nmap UDP [{target}]",
                         screenshot_name=f"phase2_nmap_udp_{safe_t}", timeout=1800)
            else:
                log(f"  — nmap UDP skipped for {target}", "dim")

        results["open_ports"] = open_ports

        if tools.get("gowitness", True) and shutil.which("gowitness") and gowitness_urls:
            log("⟶ gowitness — screenshotting web ports...", "info")
            gw_dir = p2 / "gowitness"; gw_dir.mkdir(exist_ok=True)
            unique_urls = list(dict.fromkeys(gowitness_urls))
            urls_file = gw_dir / "urls.txt"
            urls_file.write_text("\n".join(unique_urls))
            try:
                with open(gw_dir / "gowitness.log", "w") as gwl:
                    subprocess.run(
                        ["gowitness", "file", "-f", str(urls_file),
                         "-P", str(gw_dir / "screenshots"),
                         "--threads", "2",
                         "--delay", "3",
                         "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"],
                        stdout=gwl, stderr=subprocess.STDOUT,
                        text=True, timeout=900)
                log("✓ gowitness complete", "success")
                screenshots_src = gw_dir / "screenshots"
                if screenshots_src.exists():
                    for img in screenshots_src.glob("*.png"):
                        shutil.copy2(str(img), str(evidence_dir / f"gowitness_{img.name}"))
            except Exception as ex:
                log(f"✗ gowitness error: {ex}", "error")
        elif not tools.get("gowitness", True):
            log("  — gowitness skipped", "dim")
        elif not shutil.which("gowitness"):
            log("⚠ gowitness not found — skipping screenshots", "warn")

    if phases.get("vuln"):
        phase("PHASE 3 — Vulnerability Scanning")
        p3 = folder / "3_vuln"; p3.mkdir(exist_ok=True)
        vulnerabilities = {}
        missing_headers = {}
        ffuf_findings   = {}
        msf_findings    = []

        for url in url_list:
            target_key = re.sub(r"^https?://", "", url).rstrip("/")
            safe_t     = re.sub(r"[^\w\-]", "_", target_key)
            vulnerabilities[target_key] = []
            log(f"  Vuln target: {url}", "dim")

            host = re.sub(r"^https?://", "", url).rstrip("/")

            if tools.get("sslscan", True):
                run_tool(["sslscan", "--show-certificate", "-d", host],
                         p3 / f"sslscan_{safe_t}.txt",
                         f"sslscan [{host}]",
                         screenshot_name=f"phase3_sslscan_{safe_t}")
            else:
                log("  — sslscan skipped", "dim")

            if tools.get("shcheck", True):
                if shcheck_path:
                    run_tool(["python3", shcheck_path, url],
                             p3 / f"shcheck_{safe_t}.txt",
                             f"shcheck [{url}]",
                             screenshot_name=f"phase3_shcheck_{safe_t}")
                    sf = p3 / f"shcheck_{safe_t}.txt"
                    if sf.exists():
                        missing = re.findall(r"Missing security header:\s*(.+)", sf.read_text())
                        if missing:
                            missing_headers[target_key] = missing
                else:
                    log("⚠ shcheck.py not found — skipping", "warn")
            else:
                log("  — shcheck skipped", "dim")

            if tools.get("nikto", True):
                run_tool(["nikto", "-h", url, "-nocolor",
                          "-o", str(p3 / f"nikto_{safe_t}.xml"),
                          "-Format", "xml", "-nointeractive"],
                         p3 / f"nikto_{safe_t}.txt",
                         f"nikto [{url}]",
                         screenshot_name=f"phase3_nikto_{safe_t}", timeout=1800)
                nf = p3 / f"nikto_{safe_t}.txt"
                if nf.exists():
                    findings = re.findall(r"\+ (.+)", nf.read_text())
                    vulnerabilities[target_key].extend(findings[:20])
            else:
                log("  — nikto skipped", "dim")

            # ── WPScan — WordPress vulnerability scan (active only) ───
            if tools.get("wpscan", True):
                if shutil.which("wpscan"):
                    run_tool(["wpscan",
                              "--url",            url,
                              "--output",         str(p3 / f"wpscan_{safe_t}.txt"),
                              "--format",         "cli",
                              "--no-update",
                              "--enumerate",      "vp,vt,u,ap",
                              "--plugins-detection", "aggressive"],
                             p3 / f"wpscan_{safe_t}.txt",
                             f"wpscan [{url}]",
                             screenshot_name=f"phase3_wpscan_{safe_t}", timeout=1200)
                    wf = p3 / f"wpscan_{safe_t}.txt"
                    if wf.exists():
                        wp_findings = re.findall(
                            r"\[!\]\s*(.+)", wf.read_text())
                        vulnerabilities[target_key].extend(wp_findings[:20])
                else:
                    log("⚠ wpscan not found — skipping WordPress scan", "warn")
            else:
                log("  — wpscan skipped", "dim")

            if tools.get("nuclei", True):
                nuclei_out = p3 / f"nuclei_{safe_t}.txt"
                run_tool(["nuclei", "-target", url,
                          "-o", str(nuclei_out), "-nc"],
                         nuclei_out,
                         f"nuclei [{url}]",
                         screenshot_name=f"phase3_nuclei_{safe_t}", timeout=1800)
                if nuclei_out.exists():
                    nlines = [l.strip() for l in nuclei_out.read_text().splitlines() if l.strip()]
                    vulnerabilities[target_key].extend(nlines[:30])
            else:
                log("  — nuclei skipped", "dim")

            if tools.get("ffuf", True):
                wordlist = "/usr/share/seclists/Discovery/Web-Content/common.txt"
                if not Path(wordlist).exists():
                    wordlist = "/usr/share/wordlists/dirb/common.txt"
                if Path(wordlist).exists():
                    ffuf_txt = p3 / f"ffuf_{safe_t}.txt"
                    run_tool(["ffuf", "-u", f"{url}/FUZZ",
                              "-w", wordlist,
                              "-mc", "200,201,301,302,403",
                              "-t", "50"],
                             ffuf_txt, f"ffuf [{url}]",
                             screenshot_name=f"phase3_ffuf_{safe_t}")
                    if ffuf_txt.exists():
                        try:
                            found_200 = []
                            for line in ffuf_txt.read_text().splitlines():
                                if "[Status: 200," in line:
                                    parts = line.strip().split()
                                    if parts:
                                        found_200.append({
                                            "url": f"{url}/{parts[0].strip()}",
                                            "status": 200
                                        })
                            ffuf_findings[target_key] = found_200
                        except Exception:
                            pass
                else:
                    log(f"⚠ No wordlist found for ffuf [{url}]", "warn")
            else:
                log("  — ffuf skipped", "dim")

        if tools.get("metasploit", True):
            msf_rc = folder / "msf_scan.rc"
            if not msf_rc.exists():
                _write_msf_rc(msf_rc, scope_list)
            run_tool(["msfconsole", "-q", "-r", str(msf_rc)],
                     p3 / "metasploit.txt", "metasploit auxiliary scanners",
                     screenshot_name="phase3_metasploit", timeout=1800)
            msf_txt = p3 / "metasploit.txt"
            if msf_txt.exists():
                msf_findings = re.findall(r"\[\+\].*", msf_txt.read_text())
        else:
            log("  — metasploit skipped", "dim")

        results.update({
            "vulnerabilities": vulnerabilities,
            "missing_headers": missing_headers,
            "ffuf_findings":   ffuf_findings,
            "msf_findings":    msf_findings,
        })

    scan_state["results"].update(results)
    scan_state["completed"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if scan_state.get("cancel_requested"):
        log("", "info")
        log("╔══════════════════════════════════════╗", "phase")
        log("║   ⊘  SCAN CANCELLED                  ║", "phase")
        log("╚══════════════════════════════════════╝", "phase")
    else:
        log("", "info")
        log("╔══════════════════════════════════════╗", "phase")
        log("║   ✅  SCAN COMPLETE                  ║", "phase")
        log("╚══════════════════════════════════════╝", "phase")

    log("__SCAN_COMPLETE__", "control")
    scan_state["running"]          = False
    scan_state["cancel_requested"] = False


# ══════════════════════════════════════════════════════════════════════════
# USER MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

ROLES = ["Administrator", "Manager", "Junior Tester", "Representative"]

# Role permissions
ROLE_PERMS = {
    "Administrator":  {"scan": True,  "active": True,  "passive": True,  "admin": True,  "rep": False},
    "Manager":        {"scan": True,  "active": True,  "passive": True,  "admin": False, "rep": False},
    "Junior Tester":  {"scan": True,  "active": False, "passive": True,  "admin": False, "rep": False},
    "Representative": {"scan": False, "active": False, "passive": False, "admin": False, "rep": True},
}


def _hash_password(plaintext: str) -> str:
    return bcrypt.hashpw(plaintext.encode(), bcrypt.gensalt()).decode()


def _check_password(plaintext: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(plaintext.encode(), hashed.encode())
    except Exception:
        return False


def validate_password(password: str) -> list[str]:
    """Return list of unmet password requirements, empty if all pass."""
    errors = []
    if len(password) < 12:
        errors.append("Minimum 12 characters")
    if not re.search(r"[A-Z]", password):
        errors.append("At least one uppercase letter")
    if not re.search(r"[!@#$%^&*]", password):
        errors.append("At least one special character: ! @ # $ % ^ & *")
    return errors


def load_users() -> dict:
    if USERS_FILE.exists():
        try:
            return json.loads(USERS_FILE.read_text())
        except Exception:
            pass
    return {}


def save_users(users: dict):
    USERS_FILE.write_text(json.dumps(users, indent=2))


def get_user(username: str) -> dict | None:
    return load_users().get(username)


# ── Account request storage ───────────────────────────────────────────────

def load_requests() -> list:
    if REQUESTS_FILE.exists():
        try:
            return json.loads(REQUESTS_FILE.read_text())
        except Exception:
            pass
    return []


def save_requests(requests: list):
    REQUESTS_FILE.write_text(json.dumps(requests, indent=2))


# ── Client database ───────────────────────────────────────────────────────

def load_clients() -> dict:
    if CLIENTS_FILE.exists():
        try:
            return json.loads(CLIENTS_FILE.read_text())
        except Exception:
            pass
    return {}


def save_clients(clients: dict):
    CLIENTS_FILE.write_text(json.dumps(clients, indent=2))


def get_client(client_id: str) -> dict | None:
    return load_clients().get(client_id)


def clients_for_user(username: str, role: str) -> list:
    """Return clients visible to this user based on role."""
    clients = load_clients()
    if role == "Administrator":
        return list(clients.values())
    # Managers and Junior Testers only see assigned clients
    return [c for c in clients.values()
            if username in c.get("assigned_users", [])]


# ── Findings repository ───────────────────────────────────────────────────

SEVERITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4}

def load_findings() -> dict:
    if FINDINGS_FILE.exists():
        try:
            return json.loads(FINDINGS_FILE.read_text())
        except Exception:
            pass
    return {}


def save_findings(findings: dict):
    FINDINGS_FILE.write_text(json.dumps(findings, indent=2))


def get_finding(finding_id: str) -> dict | None:
    return load_findings().get(finding_id)


def client_summary(client: dict, history: list) -> dict:
    """Compute summary stats for a client from assessment history."""
    assessments = [h for h in history
                   if h.get("client_id") == client["client_id"]]
    total_findings = 0
    last_tested    = None
    for a in assessments:
        total_findings += a.get("findings_count", 0)
        if not last_tested or a.get("date", "") > last_tested:
            last_tested = a.get("date", "")
    return {
        "total_assessments": len(assessments),
        "last_tested":        last_tested or "Never",
        "total_findings":     total_findings,
        "assessments":        sorted(assessments,
                                     key=lambda x: x.get("date", ""),
                                     reverse=True),
    }


# ── Wiki storage ──────────────────────────────────────────────────────────

WIKI_DEFAULTS = {
    "assetfinder": {
        "title": "assetfinder",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/tomnomnom/assetfinder",
        "content": """# assetfinder

## What it does
assetfinder is a passive subdomain discovery tool written in Go by Tom Hudson (tomnomnom). It queries multiple public data sources to find subdomains associated with a target domain without sending any packets to the target itself.

## GitHub
[https://github.com/tomnomnom/assetfinder](https://github.com/tomnomnom/assetfinder)

## Command Artemis runs
```
assetfinder --subs-only <domain>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `--subs-only` | Only return subdomains, not the root domain itself |

## Data sources queried
- crt.sh (certificate transparency logs)
- Facebook Graph API
- VirusTotal
- Wayback Machine
- certspotter
- hackertarget
- threatcrowd

## How to read the output
Each line is a discovered subdomain, one per line:
```
mail.example.com
dev.example.com
api.example.com
```
Any subdomain listed is a valid DNS entry that was found in at least one public data source. Add these to your scope for further enumeration.

## Notes
- Entirely passive — zero contact with the target
- Fast, typically completes in under 60 seconds
- May return wildcard results like `*.example.com` — these indicate a wildcard DNS record
"""
    },
    "dnsenum": {
        "title": "dnsenum",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/fwaeytens/dnsenum",
        "content": """# dnsenum

## What it does
dnsenum is a DNS enumeration tool that performs comprehensive DNS reconnaissance against a target domain. It queries DNS records, attempts zone transfers, and brute-forces subdomains using a wordlist.

## GitHub
[https://github.com/fwaeytens/dnsenum](https://github.com/fwaeytens/dnsenum)

## Command Artemis runs
```
dnsenum --noreverse --nocolor <domain>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `--noreverse` | Skip reverse lookup on discovered IPs — keeps output clean and faster |
| `--nocolor` | Disable ANSI color codes so output is cleanly parseable |

## How to read the output
dnsenum outputs in sections:

**Host addresses** — A records for the root domain
```
example.com.    300    IN    A    93.184.216.34
```

**Name servers** — NS records, useful for zone transfer attempts
```
ns1.example.com.
ns2.example.com.
```

**Mail servers** — MX records
```
10 mail.example.com.
```

**Zone transfer** — if successful, dumps all DNS records for the domain. A successful zone transfer is a critical finding.

**Brute-forced subdomains** — results from the wordlist scan

## Notes
- Zone transfer attempts (AXFR) are active — they contact the target's nameservers directly
- A successful zone transfer reveals the entire DNS structure of the domain
"""
    },
    "spoofy": {
        "title": "spoofy",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/MattKeeley/Spoofy",
        "content": """# spoofy

## What it does
spoofy checks whether a domain can be spoofed via email. It inspects SPF, DMARC, and DKIM records to determine if an attacker could send email appearing to come from the target domain.

## GitHub
[https://github.com/MattKeeley/Spoofy](https://github.com/MattKeeley/Spoofy)

## Command Artemis runs
```
python3 /opt/spoofy/spoofy.py -d <domain> -o stdout
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-d` | Target domain to check |
| `-o stdout` | Output results to stdout rather than a file |

## How to read the output
spoofy outputs a verdict for each domain:

| Result | Meaning |
|--------|---------|
| `SPOOFABLE` | The domain can be spoofed — SPF or DMARC is missing or misconfigured |
| `NOT SPOOFABLE` | SPF and DMARC are properly configured |
| `PROBABLE` | Spoofing may be possible depending on mail server configuration |

**SPF checks** — looks for a `TXT` record like `v=spf1 include:... ~all` or `-all`. A `~all` (softfail) is weaker than `-all` (hardfail).

**DMARC checks** — looks for `_dmarc.example.com TXT v=DMARC1; p=reject`. A policy of `p=none` means no enforcement — effectively spoofable.

## Notes
- A SPOOFABLE result is a significant finding worth including in your report
- Combine with bbot-discovered email results to demonstrate the impact
"""
    },
    "o365spray": {
        "title": "o365spray",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/0xZDH/o365spray",
        "content": """# o365spray

## What it does
o365spray is a username enumeration and password spraying tool targeting Microsoft Office 365 and Azure AD. In Artemis it is used in passive enumeration mode only — confirming whether a domain uses Office 365 and optionally enumerating valid usernames.

## GitHub
[https://github.com/0xZDH/o365spray](https://github.com/0xZDH/o365spray)

## Command Artemis runs
```
python3 /opt/o365spray/o365spray.py --validate --domain <domain>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `--validate` | Check if the domain uses Microsoft O365/Azure AD — no credentials tested |
| `--domain` | Target domain |

## How to read the output
```
[*] Validating domain: example.com
[+] example.com: VALID
```

`VALID` means the domain is hosted on Microsoft O365/Azure AD. This confirms:
- Microsoft authentication endpoints are in use
- The domain's email and identity are managed by Microsoft
- Username enumeration or password spraying may be in scope (with authorization)

`INVALID` means the domain does not appear to use O365.

## Notes
- The `--validate` flag only checks DNS and Microsoft API endpoints — no credentials are tested
- This is considered passive reconnaissance
- If O365 is confirmed and user enumeration is in scope, o365spray supports `--enum -U userlist.txt`
"""
    },
    "metagoofil": {
        "title": "metagoofil",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/opsdisk/metagoofil",
        "content": """# metagoofil

## What it does
metagoofil searches Google for publicly accessible files associated with a target domain, downloads them, and extracts metadata. Metadata embedded in documents often contains usernames, software versions, email addresses, and internal paths that were never intended to be public.

In Artemis, metagoofil runs twice — once against the root domain, then again against up to 20 priority subdomains discovered during Phase 1 (prioritizing subdomains with names like mail, portal, admin, dev, vpn, sharepoint).

## GitHub
[https://github.com/opsdisk/metagoofil](https://github.com/opsdisk/metagoofil)

## Command Artemis runs

**Root domain:**
```
metagoofil -d <domain> -t pdf,doc,docx,xls,xlsx,ppt,pptx -n 20 -o <output_directory>
```

**Per subdomain (up to 20):**
```
metagoofil -d <subdomain> -t pdf,doc,docx,xls,xlsx,ppt,pptx -n 10 -o <output_directory>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-d` | Target domain or subdomain to search |
| `-t` | Comma-separated file types to search for |
| `-n` | Maximum number of files to download per run |
| `-o` | Output directory for downloaded files |

## File types targeted
| Type | Why it matters |
|------|---------------|
| `pdf` | Often contain author names, internal paths, software versions |
| `doc`, `docx` | Word documents with author and revision metadata |
| `xls`, `xlsx` | Spreadsheets may contain usernames, email lists, internal data |
| `ppt`, `pptx` | Presentations often reveal internal structure and personnel |

## Subdomain prioritization
Artemis scores subdomains by keyword relevance before running metagoofil:

| Priority keywords | Why |
|------------------|-----|
| `mail`, `portal`, `login` | Likely user-facing services with document exposure |
| `admin`, `helpdesk`, `it` | Internal tools often have leaked docs |
| `dev`, `stage`, `staging` | Development environments may expose source or config files |
| `sharepoint`, `files`, `docs` | Explicit document hosting |
| `finance`, `hr` | High-value document targets |

## How to read the output
metagoofil outputs metadata extracted per file:
```
[*] Searching for pdf files, with a maximum of 20
[*] Downloading file: https://example.com/report.pdf
[+] Author: John Smith
[+] Creator: Microsoft Word 2019
[+] Last Modified By: jsmith
[+] Producer: Adobe PDF Library
```

Key fields to note:
- **Author / Last Modified By** — real usernames, often matching Active Directory usernames
- **Creator** — software and version revealing what internal tools are in use
- **Internal paths** — Windows UNC paths like `C:/Users/jsmith/Documents/` reveal directory structure

## Notes
- A 3-second delay is applied between subdomain runs to avoid Google rate limiting
- Usernames found should be cross-referenced with bbot email results and used for o365spray enumeration if in scope
- Results are stored per-subdomain in the `1_recon/` folder
"""
    },
    "bbot": {
        "title": "bbot",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/blacklanternsecurity/bbot",
        "content": """# bbot

## What it does
bbot (Bighuge BLS OSINT Tool) is a comprehensive OSINT framework that chains together dozens of modules to perform automated attack surface mapping. It can operate in passive mode (safe, no target contact) or active mode (aggressive, full enumeration).

## GitHub
[https://github.com/blacklanternsecurity/bbot](https://github.com/blacklanternsecurity/bbot)

## Command Artemis runs

**Passive mode:**
```
bbot -t <domain> -f safe,passive -o <output_directory>
```

**Active mode:**
```
bbot -t <domain> -p everything -o <output_directory>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-t` | Target domain |
| `-f safe,passive` | Only run modules flagged as safe and passive — no target contact |
| `-p everything` | Run the `everything` preset — all available modules |
| `-o` | Output directory |

## Passive vs Active
| Mode | Contact with target | Modules run |
|------|-------------------|-------------|
| Passive | None | DNS lookups via public resolvers, certificate transparency, OSINT APIs |
| Active | Yes | DNS brute force, web crawling, port scanning, subdomain takeover checks |

## How to read the output
bbot outputs structured events:
```
[DNS_NAME] dev.example.com
[IP_ADDRESS] 93.184.216.34
[EMAIL_ADDRESS] admin@example.com
[FINDING] Possible subdomain takeover: dev.example.com
```

Event types:
| Type | Description |
|------|-------------|
| `DNS_NAME` | Discovered subdomain |
| `IP_ADDRESS` | Associated IP |
| `EMAIL_ADDRESS` | Discovered email |
| `FINDING` | Security finding worth investigating |
| `VULNERABILITY` | Confirmed vulnerability |

## Notes
- Active mode can generate significant noise — only use with authorization
- bbot's output directory contains structured JSON for further processing
"""
    },
    "amass": {
        "title": "amass",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://github.com/owasp-amass/amass",
        "content": """# amass

## What it does
amass is the OWASP Attack Surface Management tool for in-depth DNS enumeration and network mapping. It combines passive OSINT gathering with active brute forcing and DNS zone transfer attempts to build a comprehensive map of an organization's external attack surface.

## GitHub
[https://github.com/owasp-amass/amass](https://github.com/owasp-amass/amass)

## Command Artemis runs

**Passive mode:**
```
amass enum -passive -d <domain> -o <output_file> -timeout 20
```

**Active mode:**
```
amass enum -active -d <domain> -brute -w <wordlist> -o <output_file> -timeout 30
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `enum` | Subdomain enumeration subcommand |
| `-passive` | Passive mode — OSINT sources only, no target contact |
| `-active` | Active mode — direct DNS queries, zone transfers, brute force |
| `-d` | Target domain |
| `-brute` | Enable DNS brute forcing using the provided wordlist |
| `-w` | Path to wordlist for brute force (uses SecLists DNS wordlist) |
| `-o` | Output file |
| `-timeout` | Maximum runtime in minutes |

## Passive vs Active
| Mode | What it does |
|------|-------------|
| Passive | Queries certificate transparency, DNS aggregators, WHOIS, BGP data, search engines |
| Active | All passive sources plus direct nameserver queries, AXFR zone transfer attempts, DNS brute forcing |

## How to read the output
Output is one subdomain per line:
```
www.example.com
mail.example.com
dev-internal.example.com
api-v2.example.com
```

Subdomains with unexpected names (dev-internal, staging, vpn, admin) are highest priority — these may expose internal services.

## Notes
- Active mode with a large wordlist can run 30-60 minutes
- Zone transfer attempts (AXFR) in active mode are a significant test — a successful transfer is a critical finding
- Results are merged with assetfinder output in Artemis to build the master subdomain list
"""
    },
    "curl_sweep": {
        "title": "curl header sweep",
        "phase": "Phase 1 — Recon / OSINT",
        "github": "https://curl.se",
        "content": """# curl header sweep

## What it does
Artemis runs curl against each URL in scope to capture HTTP response headers. Security headers (or their absence) reveal configuration weaknesses and server information that can guide the vulnerability scanning phase.

## Command Artemis runs
```
curl -sk -I --max-time 10 <url>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-s` | Silent mode — suppresses progress output |
| `-k` | Allow insecure connections (self-signed certs) |
| `-I` | HEAD request only — fetches headers without downloading the body |
| `--max-time 10` | Timeout after 10 seconds |

## How to read the output
```
HTTP/2 200
server: nginx/1.18.0
x-powered-by: PHP/7.4.3
content-type: text/html; charset=UTF-8
strict-transport-security: max-age=31536000
```

**Key headers to note:**

| Header | Present | Missing |
|--------|---------|---------|
| `Strict-Transport-Security` | HTTPS enforced | HSTS not configured |
| `X-Frame-Options` | Clickjacking protected | Vulnerable to framing |
| `X-Content-Type-Options` | MIME sniffing prevented | MIME sniffing possible |
| `Content-Security-Policy` | XSS mitigated | No CSP in place |
| `Server` | Reveals software version | — |
| `X-Powered-By` | Reveals backend technology | — |

**Server and X-Powered-By** — revealing the exact version of nginx, Apache, or PHP narrows down which CVEs apply. These should ideally be suppressed.

## Notes
- shcheck (Phase 3) performs a more detailed security header analysis
- This sweep gives a quick overview before deep scanning begins
"""
    },
    "masscan": {
        "title": "masscan",
        "phase": "Phase 2 — Port Scanning",
        "github": "https://github.com/robertdavidgraham/masscan",
        "content": """# masscan

## What it does
masscan is the fastest port scanner available — capable of scanning the entire internet in under 6 minutes. In Artemis it performs the initial port discovery pass across all 65,535 TCP ports, then hands the confirmed open ports to nmap for detailed service fingerprinting.

## GitHub
[https://github.com/robertdavidgraham/masscan](https://github.com/robertdavidgraham/masscan)

## Command Artemis runs
```
masscan <target> -p 0-65535 --rate 1000 --wait 3 -oL <output_file>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-p 0-65535` | Scan all 65,535 TCP ports |
| `--rate 1000` | Send 1,000 packets per second — conservative rate for external testing |
| `--wait 3` | Wait 3 seconds after scan completes for late responses |
| `-oL` | Output in list format, one result per line |

## How to read the output
masscan list format:
```
open tcp 80 93.184.216.34 1714567890
open tcp 443 93.184.216.34 1714567890
open tcp 8443 93.184.216.34 1714567890
```

Fields: `status protocol port ip timestamp`

Only `open tcp` lines matter — Artemis parses these to build the port list for nmap.

## Pipeline
```
masscan (discovers open ports) → nmap (fingerprints those exact ports)
```
This two-stage approach is significantly faster than nmap `-p-` on large scopes because nmap only does the slow service detection work on ports already confirmed open.

## Notes
- Rate of 1000 pps is safe for external assessments — increase carefully on internal engagements
- Some cloud providers (AWS, Azure) rate-limit or block masscan — if no results come back, fall back to nmap directly
- masscan does not do service detection — it only confirms a port is open
"""
    },
    "nmap_tcp": {
        "title": "nmap TCP",
        "phase": "Phase 2 — Port Scanning",
        "github": "https://nmap.org",
        "content": """# nmap TCP

## What it does
nmap is the industry standard network scanner. In Artemis it runs after masscan has identified open ports, performing deep service version detection and running default NSE scripts against confirmed open ports only.

## GitHub / Docs
[https://nmap.org](https://nmap.org)

## Command Artemis runs

**When masscan found open ports:**
```
nmap -sV -sC --open -T4 -p <discovered_ports> <target>
```

**Fallback (masscan not available):**
```
nmap -sS -sV -sC -p- --open -T4 --min-rate 1000 <target>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-sV` | Version detection — probe open ports to determine service and version |
| `-sC` | Run default NSE scripts — equivalent to `--script=default` |
| `-sS` | SYN scan (stealth scan) — does not complete TCP handshake |
| `--open` | Only show open ports |
| `-T4` | Timing template 4 (aggressive) — faster on reliable networks |
| `-p` | Port specification — either discovered ports or `-p-` for all |
| `--min-rate 1000` | Send at least 1000 packets per second in fallback mode |

## Default NSE scripts run (`-sC`)
- `http-title` — grabs HTTP page title
- `http-server-header` — server banner
- `ssl-cert` — certificate details
- `ssh-hostkey` — SSH host key fingerprint
- `ftp-anon` — checks for anonymous FTP
- And many more depending on detected services

## How to read the output
```
PORT     STATE SERVICE    VERSION
22/tcp   open  ssh        OpenSSH 8.2p1 Ubuntu
80/tcp   open  http       nginx 1.18.0
443/tcp  open  ssl/https  nginx 1.18.0
8443/tcp open  ssl/http   Apache Tomcat 9.0
```

**Key things to note:**
- Exact software versions — look up CVEs for outdated versions
- Unexpected open ports — any port not needed for business purpose is an attack surface
- SSH on port 22 — note the version, check for known vulnerabilities
- Non-standard web ports (8080, 8443, 9090) — often admin interfaces or dev servers

## Notes
- Service version information directly feeds vulnerability research in Phase 3
- Outdated versions flagged here should be tested with nuclei templates
"""
    },
    "nmap_udp": {
        "title": "nmap UDP",
        "phase": "Phase 2 — Port Scanning",
        "github": "https://nmap.org",
        "content": """# nmap UDP

## What it does
UDP scanning discovers services running on UDP — a commonly overlooked protocol. Many critical services run on UDP including DNS, SNMP, DHCP, NTP, and TFTP. These services are often less hardened than their TCP counterparts.

## GitHub / Docs
[https://nmap.org](https://nmap.org)

## Command Artemis runs
```
nmap -sU --top-ports 200 -T4 <target>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-sU` | UDP scan mode |
| `--top-ports 200` | Scan only the 200 most commonly used UDP ports |
| `-T4` | Aggressive timing |

## How to read the output
```
PORT    STATE         SERVICE
53/udp  open          domain
161/udp open          snmp
123/udp open|filtered ntp
```

**States:**
| State | Meaning |
|-------|---------|
| `open` | Service is responding |
| `open/filtered` | No response received — port may be open or firewalled |
| `closed` | ICMP port unreachable received |

## High-value UDP ports
| Port | Service | Why it matters |
|------|---------|---------------|
| 53 | DNS | May allow zone transfers or recursive queries |
| 161 | SNMP | Community strings may expose configuration data |
| 69 | TFTP | Often unauthenticated file access |
| 123 | NTP | Can be used for DDoS amplification |
| 500 | IKE/IPSec | VPN endpoint identification |

## Notes
- UDP scanning is slow and unreliable — `open|filtered` is the most common state
- SNMP on UDP 161 is a high-priority finding — try community string `public` manually
- UDP scanning requires root/administrator privileges
"""
    },
    "gowitness": {
        "title": "gowitness",
        "phase": "Phase 2 — Port Scanning",
        "github": "https://github.com/sensepost/gowitness",
        "content": """# gowitness

## What it does
gowitness is a web screenshot tool that uses a headless Chrome browser to capture screenshots of every web-accessible port discovered during scanning. This gives you a visual inventory of every web service running on the target without manually visiting each URL.

## GitHub
[https://github.com/sensepost/gowitness](https://github.com/sensepost/gowitness)

## Command Artemis runs
```
gowitness file -f <urls_file> -P <screenshots_directory> --threads 5
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `file` | Read URLs from a file rather than a single target |
| `-f` | Path to file containing one URL per line |
| `-P` | Directory to save screenshots |
| `--threads 5` | Run 5 concurrent screenshot workers |

## How to use the output
Screenshots are saved as PNG files named after the URL. Open the screenshots directory and review each one:

**What to look for:**
- Login pages — admin panels, application logins, VPN portals
- Default pages — "It works!", default Apache/nginx pages indicate unfinished setup
- Error pages — may reveal software versions or internal paths
- Forgotten applications — development environments, staging servers, old admin tools
- Interesting content — anything that doesn't belong on an external-facing server

## Notes
- Artemis copies screenshots to the `Assessment_Evidence/` folder for inclusion in reports
- Any login page found should be added to the URL list for Phase 3 scanning
- Default credential testing on discovered login pages is a manual step
"""
    },
    "sslscan": {
        "title": "sslscan",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/rbsec/sslscan",
        "content": """# sslscan

## What it does
sslscan tests the SSL/TLS configuration of a web server, identifying weak cipher suites, deprecated protocol versions, certificate issues, and known vulnerabilities like POODLE, BEAST, and Heartbleed.

## GitHub
[https://github.com/rbsec/sslscan](https://github.com/rbsec/sslscan)

## Command Artemis runs
```
sslscan --show-certificate <host>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `--show-certificate` | Display the full certificate details in the output |

## How to read the output

**Protocol support:**
```
SSLv2   disabled
SSLv3   disabled
TLSv1.0 enabled
TLSv1.1 enabled
TLSv1.2 enabled
TLSv1.3 enabled
```
TLSv1.0 and TLSv1.1 are deprecated and should not be enabled. This is a medium severity finding.

**Cipher suites:**
```
Preferred TLSv1.3 128 bits TLS_AES_128_GCM_SHA256
Accepted  TLSv1.2 128 bits ECDHE-RSA-AES128-GCM-SHA256
Accepted  TLSv1.2  56 bits EXP-RC4-MD5               << EXPORT CIPHER
```
Export ciphers (56-bit) and NULL ciphers are critical findings. RC4 is deprecated.

**Certificate:**
```
Subject: CN=example.com
Issuer:  CN=Let's Encrypt Authority X3
Expires: 2025-06-01
```
Check: expiry date, whether the CN matches the domain, and whether the issuer is trusted.

## Key findings to report
| Issue | Severity |
|-------|----------|
| SSLv2 / SSLv3 enabled | Critical |
| TLS 1.0 / 1.1 enabled | Medium |
| Export or NULL ciphers | Critical |
| Self-signed certificate | Medium |
| Expired certificate | High |
| Weak key (< 2048 bit RSA) | High |
"""
    },
    "shcheck": {
        "title": "shcheck",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/santoru/shcheck",
        "content": """# shcheck

## What it does
shcheck (Security Header Check) analyzes HTTP response headers to identify missing or misconfigured security headers. Missing security headers are a common finding in web application assessments.

## GitHub
[https://github.com/santoru/shcheck](https://github.com/santoru/shcheck)

## Command Artemis runs
```
python3 /opt/shcheck/shcheck.py <url>
```

## How to read the output
```
[*] Checking security headers for: https://example.com
[!] Missing security header: Strict-Transport-Security
[!] Missing security header: X-Frame-Options
[+] Header X-Content-Type-Options is present
[!] Missing security header: Content-Security-Policy
```

## Security headers checked
| Header | Purpose | Missing = Risk |
|--------|---------|---------------|
| `Strict-Transport-Security` | Forces HTTPS | SSL stripping possible |
| `X-Frame-Options` | Prevents clickjacking | Clickjacking attacks |
| `X-Content-Type-Options` | Prevents MIME sniffing | Content injection |
| `Content-Security-Policy` | Restricts resource loading | XSS attacks |
| `Referrer-Policy` | Controls referrer information | Information leakage |
| `Permissions-Policy` | Controls browser features | Feature abuse |
| `X-XSS-Protection` | Legacy XSS protection | XSS (older browsers) |

## Severity guidance
| Finding | Severity |
|---------|----------|
| Missing CSP | Medium |
| Missing HSTS | Medium |
| Missing X-Frame-Options | Low-Medium |
| Missing X-Content-Type-Options | Low |

## Notes
- Missing security headers alone are generally low-medium severity
- Their absence is more significant when combined with other findings (e.g., missing CSP + XSS vulnerability = high)
"""
    },
    "nikto": {
        "title": "nikto",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/sullo/nikto",
        "content": """# nikto

## What it does
nikto is a web server scanner that tests for thousands of potentially dangerous files, outdated software versions, and server configuration issues. It is intentionally noisy and will be detected by most IDS/WAF systems.

## GitHub
[https://github.com/sullo/nikto](https://github.com/sullo/nikto)

## Command Artemis runs
```
nikto -h <url>
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-h` | Target host or URL |

## How to read the output
```
+ Server: Apache/2.4.29 (Ubuntu)
+ Retrieved x-powered-by header: PHP/7.2.24
+ /admin/: Admin login page/section found
+ OSVDB-3268: /icons/: Directory indexing found
+ /phpinfo.php: PHP info file found
+ Cookie PHPSESSID created without HttpOnly flag
```

Each finding starts with `+`. Key finding types:

| Finding type | Description |
|-------------|-------------|
| Outdated software | Version is known vulnerable — check CVEs |
| Interesting files | `/admin`, `/backup`, `phpinfo.php` — investigate manually |
| Directory indexing | Files are browseable — check for sensitive content |
| Cookie flags | Missing HttpOnly or Secure flags on session cookies |
| Default files | Test files, documentation left on the server |

## Notes
- nikto generates significant log entries — the target will likely detect this scan
- False positives are common — manually verify every finding before reporting
- nikto is a breadth scanner, not depth — it finds leads for manual testing
"""
    },
    "nuclei": {
        "title": "nuclei",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/projectdiscovery/nuclei",
        "content": """# nuclei

## What it does
nuclei is a template-based vulnerability scanner from ProjectDiscovery. It runs community-maintained templates that test for specific CVEs, misconfigurations, exposed panels, default credentials, and more. Templates are updated regularly to cover newly disclosed vulnerabilities.

## GitHub
[https://github.com/projectdiscovery/nuclei](https://github.com/projectdiscovery/nuclei)

## Command Artemis runs

**Passive (Phase 1):**
```
nuclei -target <url> -passive -o <output_file> -nc
```

**Active (Phase 3):**
```
nuclei -target <url> -o <output_file> -nc
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-target` | Single URL target |
| `-passive` | Passive mode — read-only requests only (Phase 1) |
| `-o` | Write results to output file |
| `-nc` | No colour — clean output with no ANSI escape codes |

## Severity levels
| Level | Examples |
|-------|---------|
| `critical` | RCE, SQLi, authentication bypass, hardcoded credentials |
| `high` | SSRF, XXE, insecure deserialization, exposed admin panels |
| `medium` | Reflected XSS, open redirect, outdated software with known exploits |
| `low` | Missing headers, information disclosure, outdated libraries |

## How to read the output
```
[CVE-2021-44228] [http] [critical] https://example.com [log4j-rce]
[exposed-panel] [http] [medium] https://example.com/admin [Admin Panel Exposed]
[default-login] [http] [high] https://example.com/manager [Apache Tomcat Default Login]
```

Format: `[template-id] [protocol] [severity] [url] [template-name]`

## Notes
- Nuclei templates are updated frequently — Artemis runs `-update-templates` on startup
- Critical findings should be manually verified before reporting
- Some templates may cause false positives on WAF-protected targets
- The template library covers thousands of CVEs and misconfigurations
"""
    },
    "ffuf": {
        "title": "ffuf",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/ffuf/ffuf",
        "content": """# ffuf

## What it does
ffuf (Fuzz Faster U Fool) is a web fuzzer used for content discovery — finding hidden directories, files, and endpoints on a web server that are not linked from public pages. It uses a wordlist to brute-force paths.

## GitHub
[https://github.com/ffuf/ffuf](https://github.com/ffuf/ffuf)

## Command Artemis runs
```
ffuf -u <url>/FUZZ -w /usr/share/seclists/Discovery/Web-Content/common.txt -mc 200,301,302,403
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `-u` | Target URL with `FUZZ` as the injection point |
| `-w` | Wordlist path — uses SecLists common.txt |
| `-mc` | Match HTTP response codes — only show these status codes |

## HTTP response codes
| Code | Meaning |
|------|---------|
| `200` | Found — page exists and is accessible |
| `301/302` | Redirect — endpoint exists, follow the redirect |
| `403` | Forbidden — endpoint exists but access is denied (still a finding) |
| `401` | Unauthorized — endpoint requires authentication |

## How to read the output
```
[Status: 200, Size: 4521, Words: 234, Lines: 89] /admin
[Status: 301, Size: 0, Words: 0, Lines: 0] /uploads
[Status: 403, Size: 287, Words: 22, Lines: 12] /.git
```

**High-priority findings:**
- `/admin`, `/administrator`, `/wp-admin` — administrative interfaces
- `/.git` — exposed git repository (critical — source code disclosure)
- `/backup`, `/old`, `/archive` — may contain sensitive data
- `/api`, `/api/v1`, `/api/v2` — API endpoints for further testing

## Notes
- 403 responses are still worth investigating — WAF or auth bypass may be possible
- An exposed `/.git` directory allows source code reconstruction — always report this as critical
- Artemis uses the SecLists `common.txt` wordlist — a broader wordlist will find more but take longer
"""
    },
    "wpscan": {
        "title": "wpscan",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/wpscanteam/wpscan",
        "content": """# wpscan

## What it does
wpscan is the WordPress security scanner. It identifies the WordPress version, installed themes and plugins, user accounts, and known vulnerabilities in any of the above. It only runs against URLs where WordPress is detected.

## GitHub
[https://github.com/wpscanteam/wpscan](https://github.com/wpscanteam/wpscan)

## Command Artemis runs
```
wpscan --url <url> --enumerate vp,vt,u,ap --plugins-detection aggressive
```

## Flag breakdown
| Flag | Description |
|------|-------------|
| `--url` | Target WordPress URL |
| `--enumerate vp` | Enumerate vulnerable plugins |
| `--enumerate vt` | Enumerate vulnerable themes |
| `--enumerate u` | Enumerate usernames |
| `--enumerate ap` | Enumerate all plugins (not just vulnerable ones) |
| `--plugins-detection aggressive` | Check all known plugin paths — slower but more thorough |

## How to read the output

**WordPress version:**
```
[+] WordPress version 6.1.1 identified (Outdated, released on 2022-11-15)
```
Note the version — check if it has known CVEs.

**Plugins:**
```
[+] akismet
   | Version: 5.0.1 (Outdated)
   | [!] 2 vulnerabilities identified:
   |     [!] CVE-2022-XXXX - XSS in admin panel
```

**Users enumerated:**
```
[+] admin
   | Found By: Author Posts - Display Name (Passive Detection)
```
Usernames are valuable for password spray or brute force (if in scope).

**Themes:**
```
[+] twentytwentythree
   | Version: 1.0 (Outdated)
```

## Notes
- wpscan requires a free API token for vulnerability data — without it, versions are identified but CVE data is not shown
- Username enumeration is active — WordPress user enumeration is a known weakness
- Only runs against confirmed WordPress installations
"""
    },
    "metasploit": {
        "title": "metasploit",
        "phase": "Phase 3 — Vulnerability Scanning",
        "github": "https://github.com/rapid7/metasploit-framework",
        "content": """# metasploit

## What it does
In Artemis, Metasploit Framework is used in auxiliary scanner mode — not for exploitation. It runs service identification and version detection modules across discovered hosts to supplement nmap results and identify services that may have known exploitable vulnerabilities.

## GitHub
[https://github.com/rapid7/metasploit-framework](https://github.com/rapid7/metasploit-framework)

## Command Artemis runs
Artemis generates a resource script (`.rc` file) and passes it to msfconsole:
```
msfconsole -r <msf_scan.rc> -q
```

The resource script contains:
```
setg RHOSTS <scope_list>
setg THREADS 10
use auxiliary/scanner/http/http_version
run
use auxiliary/scanner/ftp/ftp_version
run
use auxiliary/scanner/ssh/ssh_version
run
use auxiliary/scanner/smb/smb_version
run
```

## Modules run
| Module | What it detects |
|--------|----------------|
| `scanner/http/http_version` | Web server version and headers |
| `scanner/ftp/ftp_version` | FTP server banner and version |
| `scanner/ssh/ssh_version` | SSH server version |
| `scanner/smb/smb_version` | SMB version, OS, and hostname |

## How to read the output
```
[+] 192.168.1.100:80 - Apache httpd 2.4.29
[+] 192.168.1.100:21 - FTP Banner: 220 vsftpd 3.0.3
[+] 192.168.1.100:22 - SSH-2.0-OpenSSH_7.6p1
[+] 192.168.1.100:445 - Host is running Windows 10 (build:19041)
```

Cross-reference these versions with CVE databases to identify known vulnerabilities. Outdated service versions are direct leads for exploitation (outside Artemis scope).

## Notes
- Artemis uses Metasploit for scanning only — no exploitation modules are run
- SMB version detection revealing Windows version is useful for patch level assessment
- FTP anonymous login is tested separately if FTP is discovered
"""
    },
}


def load_wiki() -> dict:
    if WIKI_FILE.exists():
        try:
            return json.loads(WIKI_FILE.read_text())
        except Exception:
            pass
    return {}


def save_wiki(wiki: dict):
    WIKI_FILE.write_text(json.dumps(wiki, indent=2))


def get_or_init_wiki() -> dict:
    """Load wiki, seeding with defaults for any missing articles."""
    wiki = load_wiki()
    changed = False
    for slug, article in WIKI_DEFAULTS.items():
        if slug not in wiki:
            wiki[slug] = {
                "title":       article["title"],
                "phase":       article["phase"],
                "github":      article["github"],
                "content":     article["content"],
                "last_edited": None,
                "edited_by":   None,
            }
            changed = True
    if changed:
        save_wiki(wiki)
    return wiki


def bootstrap_admin():
    """
    If no users exist, create the default artemis admin account
    with a generated password and print it to the logs.
    """
    users = load_users()
    if users:
        return

    # Generate a compliant temporary password
    special  = "!@#$%^&*"
    upper    = string.ascii_uppercase
    lower    = string.ascii_lowercase
    digits   = string.digits
    # Guarantee all requirements are met
    pwd = (
        secrets.choice(upper) +
        secrets.choice(special) +
        secrets.choice(digits) +
        "".join(secrets.choice(upper + lower + digits + special)
                for _ in range(12))
    )
    # Shuffle
    pwd_list = list(pwd)
    secrets.SystemRandom().shuffle(pwd_list)
    pwd = "".join(pwd_list)

    users["artemis"] = {
        "password":   _hash_password(pwd),
        "role":       "Administrator",
        "created_at": datetime.now().isoformat(),
        "must_change": True,
    }
    save_users(users)

    banner = (
        "\n" + "="*60 +
        "\n  ARTEMIS FIRST-RUN SETUP" +
        "\n  Default admin account created:" +
        f"\n    Username : artemis" +
        f"\n    Password : {pwd}" +
        "\n  Change this password immediately after first login." +
        "\n" + "="*60 + "\n"
    )
    print(banner)


# ══════════════════════════════════════════════════════════════════════════
# AUTH DECORATORS
# ══════════════════════════════════════════════════════════════════════════

def generate_csrf_token() -> str:
    """Generate or retrieve a CSRF token for the current session."""
    try:
        if "_csrf_token" not in session:
            session["_csrf_token"] = secrets.token_hex(32)
        return session["_csrf_token"]
    except RuntimeError:
        return ""  # Outside request context


def validate_csrf(token: str) -> bool:
    """Validate a submitted CSRF token against the session token."""
    try:
        return secrets.compare_digest(
            session.get("_csrf_token", ""),
            token or ""
        )
    except RuntimeError:
        return False


# Make csrf_token available in all templates
app.jinja_env.globals["csrf_token"] = generate_csrf_token


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        # Check session exists
        if not session.get("logged_in"):
            if request.path.startswith("/api/"):
                return jsonify({"ok": False, "error": "Not authenticated"}), 401
            return redirect(url_for("login"))

        # Check session expiry (3 hours)
        login_time = session.get("login_time")
        if login_time:
            elapsed = datetime.now() - datetime.fromisoformat(login_time)
            if elapsed > timedelta(hours=SESSION_LIFETIME_HOURS):
                session.clear()
                if request.path.startswith("/api/"):
                    return jsonify({"ok": False, "error": "Session expired"}), 401
                return redirect(url_for("login", expired=1))

        # Check account expiry and active status
        username = session.get("username", "")
        user     = get_user(username)
        if user:
            # Account disabled/deactivated
            if not user.get("active", True):
                session.clear()
                if request.path.startswith("/api/"):
                    return jsonify({"ok": False, "error": "Account deactivated"}), 403
                return redirect(url_for("login", locked=1))

            # Account expiry (only applies to accounts with expires_at set)
            expires_at = user.get("expires_at")
            if expires_at:
                if datetime.now() > datetime.fromisoformat(expires_at):
                    session.clear()
                    if request.path.startswith("/api/"):
                        return jsonify({"ok": False, "error": "Account expired"}), 403
                    return redirect(url_for("login", locked=1))

        return f(*args, **kwargs)
    return decorated


def require_role(*roles):
    """Decorator to restrict a route to specific roles."""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            user_role = session.get("role", "")
            if user_role not in roles:
                if request.path.startswith("/api/"):
                    return jsonify({"ok": False,
                                    "error": f"Requires role: {' or '.join(roles)}"}), 403
                return redirect(url_for("dashboard"))
            return f(*args, **kwargs)
        return decorated
    return decorator


def admin_required(f):
    return require_role("Administrator")(f)


def active_scan_required(f):
    """Block Junior Testers from active scanning tools."""
    @wraps(f)
    def decorated(*args, **kwargs):
        role = session.get("role", "")
        if role == "Junior Tester":
            return jsonify({"ok": False,
                            "error": "Junior Testers are restricted to passive scanning."}), 403
        return f(*args, **kwargs)
    return decorated


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — AUTH
# ══════════════════════════════════════════════════════════════════════════

@app.route("/login", methods=["GET", "POST"])
def login():
    expired = request.args.get("expired")
    locked  = request.args.get("locked")
    error   = None

    if request.method == "POST":
        # CSRF check
        if not validate_csrf(request.form.get("_csrf_token", "")):
            audit_log("CSRF_BLOCKED", request.form.get("username","—"))
            error = "Invalid request — please try again."
            return render_template("login.html", error=error,
                                   expired=expired, locked=locked)

        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        ip       = request.remote_addr

        # ── Rate limiting check ─────────────────────────────────────────
        now  = _time.time()
        info = _login_attempts.get(ip, {"count": 0, "lockout_until": 0})
        if info["lockout_until"] > now:
            remaining = int(info["lockout_until"] - now)
            audit_log("LOGIN_BLOCKED", username, f"IP locked out — {remaining}s remaining")
            error = f"Too many failed attempts. Try again in {remaining // 60 + 1} minute(s)."
            return render_template("login.html", error=error,
                                   expired=expired, locked=locked)

        user = get_user(username)
        if user and _check_password(password, user["password"]):
            # Success — clear attempts
            _login_attempts.pop(ip, None)
            session.clear()
            session["logged_in"]  = True
            session["username"]   = username
            session["role"]       = user["role"]
            session["login_time"] = datetime.now().isoformat()
            session.permanent     = False

            audit_log("LOGIN_SUCCESS", username)

            if user.get("must_change"):
                return redirect(url_for("change_password", first=1))
            if user["role"] == "Representative":
                return redirect(url_for("rep_portal"))
            return redirect(url_for("dashboard"))

        # Failed attempt
        info["count"] += 1
        if info["count"] >= LOGIN_MAX_ATTEMPTS:
            info["lockout_until"] = now + LOGIN_LOCKOUT_SECS
            info["count"]         = 0
            audit_log("LOGIN_LOCKOUT", username, f"Locked after {LOGIN_MAX_ATTEMPTS} failures")
            error = f"Too many failed attempts. Account locked for {LOGIN_LOCKOUT_SECS // 60} minutes."
        else:
            remaining_attempts = LOGIN_MAX_ATTEMPTS - info["count"]
            audit_log("LOGIN_FAILED", username, f"Attempt {info['count']}/{LOGIN_MAX_ATTEMPTS}")
            error = f"Invalid credentials. {remaining_attempts} attempt(s) remaining."
        _login_attempts[ip] = info

    return render_template("login.html", error=error, expired=expired, locked=locked)


@app.route("/logout")
def logout():
    audit_log("LOGOUT", session.get("username", "—"))
    session.clear()
    return redirect(url_for("login"))


@app.route("/change-password", methods=["GET", "POST"])
@login_required
def change_password():
    first  = request.args.get("first")
    errors = []

    if request.method == "POST":
        current  = request.form.get("current_password", "")
        new_pwd  = request.form.get("new_password", "")
        confirm  = request.form.get("confirm_password", "")
        username = session.get("username")
        user     = get_user(username)

        if not _check_password(current, user["password"]):
            errors.append("Current password is incorrect.")
        elif new_pwd != confirm:
            errors.append("New passwords do not match.")
        else:
            errors = validate_password(new_pwd)

        if not errors:
            users = load_users()
            users[username]["password"]    = _hash_password(new_pwd)
            users[username]["must_change"] = False
            save_users(users)
            return redirect(url_for("dashboard"))

    return render_template("change_password.html", errors=errors, first=first)


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — ADMIN USER MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@app.route("/admin/audit")
@login_required
@admin_required
def audit_log_view():
    lines = []
    if AUDIT_LOG_FILE.exists():
        raw = AUDIT_LOG_FILE.read_text(errors="ignore").splitlines()
        lines = list(reversed(raw))  # most recent first
    return render_template("audit_log.html",
                           lines=lines,
                           role=session.get("role", ""),
                           username=session.get("username", ""))


@app.route("/admin/users")
@login_required
@admin_required
def admin_users():
    users = load_users()
    # Strip password hashes before passing to template
    safe_users = {
        u: {k: v for k, v in data.items() if k != "password"}
        for u, data in users.items()
    }
    return render_template("admin_users.html",
                           users=safe_users, roles=ROLES,
                           current_user=session.get("username"),
                           now=datetime.now().isoformat())


@app.route("/admin/users/create", methods=["POST"])
@login_required
@admin_required
def admin_create_user():
    data     = request.get_json()
    username = data.get("username", "").strip().lower()
    password = data.get("password", "")
    role     = data.get("role", "Junior Tester")

    if not username:
        return jsonify({"ok": False, "error": "Username is required."}), 400
    if role not in ROLES:
        return jsonify({"ok": False, "error": "Invalid role."}), 400

    users = load_users()
    if username in users:
        return jsonify({"ok": False, "error": "Username already exists."}), 400

    pwd_errors = validate_password(password)
    if pwd_errors:
        return jsonify({"ok": False, "error": "; ".join(pwd_errors)}), 400

    users[username] = {
        "password":    _hash_password(password),
        "role":        role,
        "created_at":  datetime.now().isoformat(),
        "must_change": False,
    }
    save_users(users)
    return jsonify({"ok": True})


@app.route("/admin/users/update", methods=["POST"])
@login_required
@admin_required
def admin_update_user():
    data     = request.get_json()
    username = data.get("username", "").strip().lower()
    role     = data.get("role")
    password = data.get("password", "").strip()

    users = load_users()
    if username not in users:
        return jsonify({"ok": False, "error": "User not found."}), 404

    if role:
        if role not in ROLES:
            return jsonify({"ok": False, "error": "Invalid role."}), 400
        # Prevent removing the last admin
        if users[username]["role"] == "Administrator" and role != "Administrator":
            admins = [u for u, d in users.items() if d["role"] == "Administrator"]
            if len(admins) <= 1:
                return jsonify({"ok": False,
                                "error": "Cannot demote the last Administrator."}), 400
        users[username]["role"] = role

    if password:
        pwd_errors = validate_password(password)
        if pwd_errors:
            return jsonify({"ok": False, "error": "; ".join(pwd_errors)}), 400
        users[username]["password"]    = _hash_password(password)
        users[username]["must_change"] = True

    save_users(users)
    return jsonify({"ok": True})


@app.route("/admin/users/reactivate", methods=["POST"])
@login_required
@admin_required
def admin_reactivate_user():
    data     = request.get_json()
    username = data.get("username", "").strip().lower()
    days     = int(data.get("days", 14))

    users = load_users()
    if username not in users:
        return jsonify({"ok": False, "error": "User not found."}), 404

    users[username]["active"]     = True
    users[username]["expires_at"] = (datetime.now() + timedelta(days=days)).isoformat()
    save_users(users)
    return jsonify({"ok": True,
                    "expires_at": users[username]["expires_at"]})


@app.route("/admin/users/delete", methods=["POST"])
@login_required
@admin_required
def admin_delete_user():
    data     = request.get_json()
    username = data.get("username", "").strip().lower()

    if username == session.get("username"):
        return jsonify({"ok": False, "error": "Cannot delete your own account."}), 400

    users = load_users()
    if username not in users:
        return jsonify({"ok": False, "error": "User not found."}), 404

    # Prevent deleting the last admin
    if users[username]["role"] == "Administrator":
        admins = [u for u, d in users.items() if d["role"] == "Administrator"]
        if len(admins) <= 1:
            return jsonify({"ok": False,
                            "error": "Cannot delete the last Administrator."}), 400

    del users[username]
    save_users(users)
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════

@app.route("/")
@login_required
def dashboard():
    # Representatives have their own portal
    if session.get("role") == "Representative":
        return redirect(url_for("rep_portal"))
    username = session.get("username", "")
    role     = session.get("role", "")
    history  = load_history_for_user(username)[:5]
    clients  = clients_for_user(username, role)
    return render_template("dashboard.html",
                           history=history,
                           clients=clients,
                           username=username,
                           role=role)


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — SCAN
# ══════════════════════════════════════════════════════════════════════════

@app.route("/scan")
@login_required
def scan_page():
    role     = session.get("role", "")
    username = session.get("username", "")
    clients  = clients_for_user(username, role)
    return render_template("index.html",
                           role=role,
                           username=username,
                           clients=clients)


@app.route("/api/startup-status")
@login_required
def startup_status():
    tp = scan_state.get("tool_paths", {})
    return jsonify({
        "shcheck":  bool(tp.get("shcheck")),
        "spoofy":   bool(tp.get("spoofy")),
        "o365scan": bool(tp.get("o365scan")),
        "ready":    True,
    })


@app.route("/api/submit", methods=["POST"])
@login_required
def submit():
    if scan_state["running"]:
        return jsonify({"ok": False, "error": "Scan already in progress."}), 409

    data      = request.get_json()
    client    = data.get("client", "").strip()
    date      = data.get("date", datetime.now().strftime("%Y-%m-%d")).strip()
    domain    = re.sub(r"^https?://", "", data.get("domain", "").strip()).rstrip("/").lower()
    scope_raw = data.get("scope", "").strip()
    urls_raw  = data.get("urls", "").strip()

    if not client:
        return jsonify({"ok": False, "error": "Client name is required."}), 400
    if not scope_raw and not urls_raw and not domain:
        return jsonify({"ok": False, "error": "Enter at least one target."}), 400

    scope_list = [re.sub(r"^https?://", "", l).rstrip("/").strip()
                  for l in scope_raw.splitlines() if l.strip()]
    url_list   = [l.strip() for l in urls_raw.splitlines() if l.strip()]

    client_id  = data.get("client_id", "").strip()

    # Generate a unique scan ID for the folder name
    import uuid
    scan_id    = str(uuid.uuid4())

    # Determine parent folder — client UUID dir or Quick_Scans
    if client_id:
        clients = load_clients()
        if client_id in clients:
            parent = RESULTS_BASE / client_id
        else:
            parent = QUICK_SCANS_BASE
    else:
        parent = QUICK_SCANS_BASE

    parent.mkdir(parents=True, exist_ok=True)
    folder = parent / scan_id
    folder.mkdir(parents=True, exist_ok=True)

    # Store a metadata file inside the scan folder
    meta = {
        "client":    client,
        "client_id": client_id,
        "date":      date,
        "domain":    domain,
        "scan_id":   scan_id,
        "created_at": datetime.now().isoformat(),
    }
    (folder / "meta.json").write_text(json.dumps(meta, indent=2))

    if scope_list: (folder / "scope.txt").write_text("\n".join(scope_list))
    if url_list:   (folder / "urls.txt").write_text("\n".join(url_list))
    if domain:     (folder / "domain.txt").write_text(domain)

    scan_state["client_folder"] = folder
    scan_state["results"] = {
        "client":    client,
        "client_id": client_id,
        "date":      date,
        "domain":    domain,
        "scope_list": scope_list,
        "url_list":   url_list,
        "scan_id":    scan_id,
    }

    return jsonify({
        "ok":          True,
        "scan_id":     scan_id,
        "scope_count": len(scope_list),
        "url_count":   len(url_list),
        "domain":      domain,
    })


@app.route("/api/start", methods=["POST"])
@login_required
def start_scan():
    if scan_state["running"]:
        return jsonify({"ok": False, "error": "Scan already running."}), 409

    folder = scan_state.get("client_folder")
    if not folder:
        return jsonify({"ok": False, "error": "Submit engagement details first."}), 400

    data = request.get_json()

    # Phase-level toggles
    phases = {
        "recon": data.get("recon", True),
        "scan":  data.get("scan",  True),
        "vuln":  data.get("vuln",  True),
    }

    # Per-tool toggles (default True so old clients still work)
    tools = {
        "assetfinder":  data.get("assetfinder",  True),
        "dnsenum":      data.get("dnsenum",       True),
        "spoofy":       data.get("spoofy",        True),
        "o365spray":    data.get("o365spray",     True),
        "metagoofil":   data.get("metagoofil",    True),
        "bbot":         data.get("bbot",          True),
        "bbot_mode":    data.get("bbot_mode",     "passive"),
        "amass":        data.get("amass",         True),
        "amass_mode":   data.get("amass_mode",    "passive"),
        "nuclei_passive": data.get("nuclei_passive", True),
        "breach_check":   data.get("breach_check",   True),
        "curl_sweep":   data.get("curl_sweep",    True),
        "masscan":      data.get("masscan",       True),
        "nmap_tcp":     data.get("nmap_tcp",      True),
        "nmap_udp":     data.get("nmap_udp",      True),
        "gowitness":    data.get("gowitness",     True),
        "sslscan":      data.get("sslscan",       True),
        "shcheck":      data.get("shcheck",       True),
        "nikto":        data.get("nikto",         True),
        "nuclei":       data.get("nuclei",        True),
        "ffuf":         data.get("ffuf",          True),
        "wpscan":       data.get("wpscan",        True),
        "metasploit":   data.get("metasploit",    True),
    }

    # ── Role enforcement — server-side, regardless of what the UI sends ──
    role = session.get("role", "")
    if role == "Junior Tester":
        # Force passive mode — disable all active tools
        ACTIVE_TOOLS = ["curl_sweep", "masscan", "nmap_tcp", "nmap_udp", "gowitness",
                        "nikto", "nuclei", "ffuf", "wpscan", "metasploit"]
        for t in ACTIVE_TOOLS:
            tools[t] = False
        tools["bbot_mode"]   = "passive"
        tools["amass_mode"]  = "passive"   # amass allowed but passive only
        phases["scan"] = False             # entire port scan phase disabled for JT

    results    = scan_state["results"]
    scope_list = results.get("scope_list", [])
    url_list   = results.get("url_list",   [])
    domain     = results.get("domain",     "")
    tool_paths = scan_state.get("tool_paths", {})

    while not scan_state["log_queue"].empty():
        scan_state["log_queue"].get_nowait()

    scan_state["running"]   = True
    scan_state["started"]   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    scan_state["completed"] = ""

    results = scan_state.get("results", {})
    audit_log("SCAN_STARTED", session.get("username", "—"),
              f"client={results.get('client','?')} domain={results.get('domain','?')}")
    t = threading.Thread(
        target=run_scan,
        args=(scope_list, url_list, domain, phases, tools, folder, tool_paths),
        daemon=True,
    )
    t.start()
    return jsonify({"ok": True})


@app.route("/api/stream")
@login_required
def stream():
    def event_generator():
        while True:
            try:
                item = scan_state["log_queue"].get(timeout=30)
                yield f"data: {item['tag']}|{item['ts']}|{item['msg']}\n\n"
                if item.get("tag") == "control" and item.get("msg") == "__SCAN_COMPLETE__":
                    break
            except queue.Empty:
                yield ": keepalive\n\n"

    return Response(event_generator(),
                    mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/report", methods=["POST"])
@login_required
def generate_report():
    folder  = scan_state.get("client_folder")
    results = scan_state.get("results", {})

    if not folder or not results.get("client"):
        return jsonify({"ok": False, "error": "No scan data available."}), 400

    try:
        from report_generator import generate_reports
        pdf_path, docx_path = generate_reports(
            results, Path(folder),
            log_fn=lambda m, t="info": scan_state["log_queue"].put(
                {"ts": datetime.now().strftime("%H:%M:%S"), "msg": m, "tag": t}
            )
        )
        return jsonify({
            "ok":   True,
            "pdf":  f"/api/download/{pdf_path.name}",
            "docx": f"/api/download/{docx_path.name}",
        })
    except Exception as ex:
        return jsonify({"ok": False, "error": str(ex)}), 500


@app.route("/api/download/<filename>")
@login_required
def download_file(filename):
    folder = scan_state.get("client_folder")
    if not folder:
        return "No active session", 404
    file_path = Path(folder) / filename
    if not file_path.exists():
        return "File not found", 404
    return send_file(str(file_path), as_attachment=True)


@app.route("/api/delete-assessment", methods=["POST"])
@login_required
def delete_assessment():
    """Remove an entry from dashboard history — files on disk are untouched."""
    data   = request.get_json()
    folder = data.get("folder", "").strip()
    if not folder:
        return jsonify({"ok": False, "error": "No folder specified."}), 400

    history = load_history()
    history = [h for h in history if h.get("folder") != folder]
    HISTORY_FILE.write_text(json.dumps(history, indent=2))
    return jsonify({"ok": True})


@app.route("/api/save-assessment", methods=["POST"])
@login_required
def save_assessment_route():
    """Called when the user clicks Exit — saves to dashboard history."""
    results = scan_state.get("results", {})
    folder  = scan_state.get("client_folder")

    if not results.get("client"):
        return jsonify({"ok": False, "error": "No assessment to save."}), 400

    data           = request.get_json() or {}
    phases         = data.get("phases", {"recon": True, "scan": True, "vuln": True})
    client_id      = data.get("client_id", "") or results.get("client_id", "")
    findings_count = sum(len(v) for v in results.get("vulnerabilities", {}).values())

    save_assessment(
        client         = results.get("client", "Unknown"),
        date           = results.get("date",   ""),
        domain         = results.get("domain", ""),
        folder         = str(folder) if folder else "",
        phases         = phases,
        username       = session.get("username", ""),
        started        = scan_state.get("started",   ""),
        completed      = scan_state.get("completed", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        client_id      = client_id,
        findings_count = findings_count,
    )

    audit_log("SCAN_SAVED", session.get("username", "—"),
              f"client={results.get('client','?')} domain={results.get('domain','?')}")

    # Reset state
    scan_state["client_folder"] = None
    scan_state["results"]       = {}
    scan_state["started"]       = ""
    scan_state["completed"]     = ""
    while not scan_state["log_queue"].empty():
        scan_state["log_queue"].get_nowait()

    return jsonify({"ok": True, "redirect": "/"})


@app.route("/api/cancel", methods=["POST"])
@login_required
def cancel_scan():
    if not scan_state["running"]:
        return jsonify({"ok": False, "error": "No scan running."}), 400
    scan_state["cancel_requested"] = True
    return jsonify({"ok": True, "message": "Cancel requested — stopping after current tool."})


@app.route("/api/reset", methods=["POST"])
@login_required
def reset():
    if scan_state["running"]:
        return jsonify({"ok": False, "error": "Scan still running."}), 409
    scan_state["client_folder"]    = None
    scan_state["results"]          = {}
    scan_state["started"]          = ""
    scan_state["completed"]        = ""
    scan_state["cancel_requested"] = False
    while not scan_state["log_queue"].empty():
        scan_state["log_queue"].get_nowait()
    return jsonify({"ok": True})


@app.route("/api/assessment/<folder_name>/meta", methods=["POST"])
@login_required
def save_assessment_meta(folder_name):
    """Save editable fields (notes, executive_summary, risk_register) to meta.json."""
    # Locate the folder by searching results tree
    target = None
    for p in RESULTS_BASE.rglob("meta.json"):
        if p.parent.name == folder_name:
            target = p
            break
    if not target:
        return jsonify({"ok": False, "error": "Assessment not found."}), 404

    try:
        meta = json.loads(target.read_text())
    except Exception:
        meta = {}

    data = request.get_json() or {}
    allowed = {"notes", "executive_summary", "risk_register"}
    for key in allowed:
        if key in data:
            meta[key] = data[key]

    target.write_text(json.dumps(meta, indent=2))
    return jsonify({"ok": True})


# ══════════════════════════════════════════════════════════════════════════
# ASSESSMENT SUMMARY
# ══════════════════════════════════════════════════════════════════════════

def _read_assessment_data(folder: Path, entry: dict) -> dict:
    """Read scan output files and build a summary data structure."""
    data = {
        "client":            entry.get("client", ""),
        "date":              entry.get("date", ""),
        "domain":            entry.get("domain", "—"),
        "folder":            str(folder),
        "started":           entry.get("started", ""),
        "completed":         entry.get("completed", ""),
        "phases":            entry.get("phases", {}),
        "scope":             [],
        "urls":              [],
        "subdomains":        [],
        "notes":             "",
        "executive_summary": "",
        "risk_register":     [],
        "open_ports":{},
        "emails":    [],
        "vuln_count":0,
        "missing_headers":{},
        "log_lines": 0,
        "duration":  "",
    }

    # Read meta.json if present (new UUID-based folder structure)
    meta_f = folder / "meta.json"
    if meta_f.exists():
        try:
            meta = json.loads(meta_f.read_text())
            data["client"]            = data["client"]    or meta.get("client", "")
            data["domain"]            = data["domain"]    or meta.get("domain", "—")
            data["client_id"]         = meta.get("client_id", "")
            data["notes"]             = meta.get("notes", "")
            data["executive_summary"] = meta.get("executive_summary", "")
            data["risk_register"]     = meta.get("risk_register", [])
        except Exception:
            pass

    # Read target files
    scope_f = folder / "scope.txt"
    if scope_f.exists():
        data["scope"] = [l.strip() for l in scope_f.read_text().splitlines() if l.strip()]

    urls_f = folder / "urls.txt"
    if urls_f.exists():
        data["urls"] = [l.strip() for l in urls_f.read_text().splitlines() if l.strip()]

    # Log line count and duration
    log_f = folder / "artemis.log"
    if log_f.exists():
        lines = log_f.read_text().splitlines()
        data["log_lines"] = len(lines)
        # Extract first and last timestamps from log
        if lines:
            try:
                first = lines[0].split("]")[0].lstrip("[")
                last  = lines[-1].split("]")[0].lstrip("[")
                data["log_start"] = first
                data["log_end"]   = last
            except Exception:
                pass

    # Calculate duration from stored timestamps
    if data["started"] and data["completed"]:
        try:
            fmt = "%Y-%m-%d %H:%M:%S"
            s   = datetime.strptime(data["started"],   fmt)
            e   = datetime.strptime(data["completed"], fmt)
            secs = int((e - s).total_seconds())
            h, r = divmod(secs, 3600)
            m, s = divmod(r, 60)
            if h:   data["duration"] = f"{h}h {m}m {s}s"
            elif m: data["duration"] = f"{m}m {s}s"
            else:   data["duration"] = f"{s}s"
        except Exception:
            pass

    # Subdomains from assetfinder / theharvester
    subs = set()
    for fname in (folder / "1_recon").glob("assetfinder_*.txt") if (folder / "1_recon").exists() else []:
        subs.update(l.strip() for l in fname.read_text().splitlines() if l.strip())
    data["subdomains"] = list(subs)

    # Open ports from nmap
    p2 = folder / "2_scan"
    if p2.exists():
        for f in p2.glob("nmap_tcp_*.txt"):
            host = f.stem.replace("nmap_tcp_", "")
            ports = re.findall(r"(\d+)/tcp\s+open\s+(\S+)", f.read_text())
            if ports:
                data["open_ports"][host] = [f"{p}/{svc}" for p, svc in ports]

    # Vuln count from nuclei + nikto
    p3 = folder / "3_vuln"
    vuln_count = 0
    if p3.exists():
        for f in p3.glob("nuclei_*.txt"):
            vuln_count += sum(1 for l in f.read_text().splitlines() if l.strip())
        for f in p3.glob("nikto_*.txt"):
            vuln_count += len(re.findall(r"\+ ", f.read_text()))
        # Missing headers
        for f in p3.glob("shcheck_*.txt"):
            host = f.stem.replace("shcheck_", "")
            missing = re.findall(r"Missing security header:\s*(.+)", f.read_text())
            if missing:
                data["missing_headers"][host] = missing
    data["vuln_count"] = vuln_count

    return data


@app.route("/assessment/<path:folder_name>")
@login_required
def assessment_summary(folder_name):
    username = session.get("username", "")
    # Find the entry in this user's history
    history = load_history_for_user(username)
    entry   = next((h for h in history if Path(h["folder"]).name == folder_name
                    or h["folder"].endswith(folder_name)), None)

    # Administrators can view any assessment
    if not entry and session.get("role") == "Administrator":
        entry = next((h for h in load_history()
                      if Path(h["folder"]).name == folder_name
                      or h["folder"].endswith(folder_name)), None)

    if not entry:
        return "Assessment not found or access denied.", 404

    folder = Path(entry["folder"])
    if not folder.exists():
        return "Assessment folder not found on disk.", 404

    data = _read_assessment_data(folder, entry)
    return render_template("assessment_summary.html", data=data,
                           role=session.get("role", ""),
                           username=username)


# ── Export routes ─────────────────────────────────────────────────────────

@app.route("/assessment/<path:folder_name>/export/zip")
@login_required
def export_zip(folder_name):
    import zipfile, io
    username = session.get("username", "")
    history  = load_history_for_user(username)
    if session.get("role") == "Administrator":
        history = load_history()
    entry = next((h for h in history if Path(h["folder"]).name == folder_name
                  or h["folder"].endswith(folder_name)), None)
    if not entry:
        return "Assessment not found.", 404

    folder = Path(entry["folder"])
    if not folder.exists():
        return "Folder not found on disk.", 404

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in folder.rglob("*"):
            if file.is_file():
                zf.write(file, file.relative_to(folder.parent))
    buf.seek(0)
    safe_name = re.sub(r"[^\w\-]", "_", entry.get("client", folder_name))
    return send_file(buf, as_attachment=True,
                     download_name=f"Artemis_{safe_name}_{entry.get('date','')}.zip",
                     mimetype="application/zip")


@app.route("/assessment/<path:folder_name>/export/txt")
@login_required
def export_txt(folder_name):
    import io
    username = session.get("username", "")
    history  = load_history_for_user(username)
    if session.get("role") == "Administrator":
        history = load_history()
    entry = next((h for h in history if Path(h["folder"]).name == folder_name
                  or h["folder"].endswith(folder_name)), None)
    if not entry:
        return "Assessment not found.", 404
    folder = Path(entry["folder"])
    if not folder.exists():
        return "Folder not found on disk.", 404

    data = _read_assessment_data(folder, entry)
    lines = [
        "=" * 60,
        f"ARTEMIS ASSESSMENT REPORT",
        "=" * 60,
        f"Client:    {data['client']}",
        f"Date:      {data['date']}",
        f"Domain:    {data['domain']}",
        f"Started:   {data['started']}",
        f"Completed: {data['completed']}",
        f"Duration:  {data['duration'] or '—'}",
        "",
        "SCOPE",
        "-" * 40,
    ] + (data["scope"] or ["—"]) + [
        "",
        "URLs",
        "-" * 40,
    ] + (data["urls"] or ["—"]) + [
        "",
        "SUBDOMAINS DISCOVERED",
        "-" * 40,
    ] + (data["subdomains"] or ["None found"]) + [
        "",
        "OPEN PORTS",
        "-" * 40,
    ]
    for host, ports in data["open_ports"].items():
        lines.append(f"{host}: {', '.join(ports)}")
    if not data["open_ports"]:
        lines.append("None found")
    lines += [
        "",
        "EMAILS FOUND",
        "-" * 40,
    ] + (list(set(data["emails"])) or ["None found"]) + [
        "",
        f"VULNERABILITIES / FINDINGS: {data['vuln_count']}",
        "",
        "MISSING SECURITY HEADERS",
        "-" * 40,
    ]
    for host, headers in data["missing_headers"].items():
        lines.append(f"{host}:")
        for h in headers:
            lines.append(f"  - {h}")
    if not data["missing_headers"]:
        lines.append("None found")
    lines += ["", "=" * 60, "Generated by Artemis // Retrobyte Cybersecurity LLC", "=" * 60]

    content = "\n".join(lines)
    buf = io.BytesIO(content.encode("utf-8"))
    safe = re.sub(r"[^\w\-]", "_", data["client"])
    return send_file(buf, as_attachment=True,
                     download_name=f"Artemis_{safe}_{data['date']}.txt",
                     mimetype="text/plain")


@app.route("/assessment/<path:folder_name>/export/docx")
@login_required
def export_docx(folder_name):
    from docx import Document
    from docx.shared import Pt, RGBColor
    import io
    username = session.get("username", "")
    history  = load_history_for_user(username)
    if session.get("role") == "Administrator":
        history = load_history()
    entry = next((h for h in history if Path(h["folder"]).name == folder_name
                  or h["folder"].endswith(folder_name)), None)
    if not entry:
        return "Assessment not found.", 404
    folder = Path(entry["folder"])
    if not folder.exists():
        return "Folder not found on disk.", 404

    data = _read_assessment_data(folder, entry)
    doc  = Document()

    # Title
    title = doc.add_heading("Artemis Assessment Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x3A, 0x5A, 0x7A)

    # Meta table
    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Client",    data["client"]),
        ("Date",      data["date"]),
        ("Domain",    data["domain"]),
        ("Started",   data["started"]),
        ("Completed", data["completed"]),
        ("Duration",  data["duration"] or "—"),
    ]
    for i, (label, value) in enumerate(fields):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = value

    def add_section(title, items):
        doc.add_heading(title, level=1)
        if items:
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("None found.")

    doc.add_paragraph()
    add_section("Scope (IPs / Hosts)", data["scope"])
    add_section("URLs Scanned", data["urls"])
    add_section("Subdomains Discovered", data["subdomains"])

    doc.add_heading("Open Ports", level=1)
    if data["open_ports"]:
        for host, ports in data["open_ports"].items():
            doc.add_paragraph(f"{host}: {', '.join(ports)}", style="List Bullet")
    else:
        doc.add_paragraph("None found.")

    add_section("Emails Found", list(set(data["emails"])))

    doc.add_heading("Missing Security Headers", level=1)
    if data["missing_headers"]:
        for host, headers in data["missing_headers"].items():
            doc.add_paragraph(host, style="List Bullet")
            for h in headers:
                doc.add_paragraph(f"  • {h}")
    else:
        doc.add_paragraph("None found.")

    doc.add_heading(f"Total Findings / Vulnerabilities: {data['vuln_count']}", level=1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = re.sub(r"[^\w\-]", "_", data["client"])
    return send_file(buf, as_attachment=True,
                     download_name=f"Artemis_{safe}_{data['date']}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/assessment/<path:folder_name>/export/xlsx")
@login_required
def export_xlsx(folder_name):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        subprocess.run(["pip3", "install", "openpyxl", "--break-system-packages"],
                       capture_output=True)
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    import io

    username = session.get("username", "")
    history  = load_history_for_user(username)
    if session.get("role") == "Administrator":
        history = load_history()
    entry = next((h for h in history if Path(h["folder"]).name == folder_name
                  or h["folder"].endswith(folder_name)), None)
    if not entry:
        return "Assessment not found.", 404
    folder = Path(entry["folder"])
    if not folder.exists():
        return "Folder not found on disk.", 404

    data = _read_assessment_data(folder, entry)
    wb   = openpyxl.Workbook()

    hdr_font  = Font(bold=True, color="FFFFFF")
    hdr_fill  = PatternFill("solid", fgColor="3A5A7A")
    hdr_align = Alignment(horizontal="center")

    def make_sheet(title, headers, rows):
        ws = wb.create_sheet(title=title)
        for ci, h in enumerate(headers, 1):
            c = ws.cell(1, ci, h)
            c.font  = hdr_font
            c.fill  = hdr_fill
            c.alignment = hdr_align
            ws.column_dimensions[c.column_letter].width = 28
        for ri, row in enumerate(rows, 2):
            for ci, val in enumerate(row, 1):
                ws.cell(ri, ci, val)
        return ws

    # Summary sheet
    ws = wb.active
    ws.title = "Summary"
    for ci, h in enumerate(["Field", "Value"], 1):
        c = ws.cell(1, ci, h)
        c.font = hdr_font; c.fill = hdr_fill
        ws.column_dimensions[c.column_letter].width = 30
    summary_rows = [
        ("Client",    data["client"]),
        ("Date",      data["date"]),
        ("Domain",    data["domain"]),
        ("Started",   data["started"]),
        ("Completed", data["completed"]),
        ("Duration",  data["duration"] or "—"),
        ("Total Findings", str(data["vuln_count"])),
        ("Subdomains Found", str(len(data["subdomains"]))),
        ("Open Port Hosts",  str(len(data["open_ports"]))),
    ]
    for ri, (f, v) in enumerate(summary_rows, 2):
        ws.cell(ri, 1, f); ws.cell(ri, 2, v)

    make_sheet("Scope",      ["Host / IP"],  [[s] for s in data["scope"]])
    make_sheet("URLs",       ["URL"],        [[u] for u in data["urls"]])
    make_sheet("Subdomains", ["Subdomain"],  [[s] for s in data["subdomains"]])
    make_sheet("Open Ports", ["Host", "Port / Service"],
               [[h, p] for h, ports in data["open_ports"].items() for p in ports])
    make_sheet("Emails",     ["Email"],      [[e] for e in set(data["emails"])])
    make_sheet("Missing Headers", ["Host", "Missing Header"],
               [[h, hdr] for h, hdrs in data["missing_headers"].items() for hdr in hdrs])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe = re.sub(r"[^\w\-]", "_", data["client"])
    return send_file(buf, as_attachment=True,
                     download_name=f"Artemis_{safe}_{data['date']}.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — REPRESENTATIVE PORTAL
# ══════════════════════════════════════════════════════════════════════════

def rep_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get("role") != "Representative":
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)
    return decorated


@app.route("/rep")
@login_required
@rep_required
def rep_portal():
    username = session.get("username", "")
    # Show only requests submitted by this rep
    all_requests = load_requests()
    my_requests  = [r for r in all_requests if r.get("submitted_by") == username]
    return render_template("rep_portal.html",
                           username=username,
                           requests=my_requests)


@app.route("/api/requests/submit", methods=["POST"])
@login_required
@rep_required
def submit_request():
    data = request.get_json()
    name      = data.get("name", "").strip()
    email     = data.get("email", "").strip().lower()
    note      = data.get("note", "").strip()
    username  = session.get("username", "")

    if not name:
        return jsonify({"ok": False, "error": "Name is required."}), 400
    if not email or "@" not in email:
        return jsonify({"ok": False, "error": "Valid email is required."}), 400

    # Role is always Junior Tester for rep-submitted requests
    role = "Junior Tester"

    # Check email not already a username
    users = load_users()
    if email in users:
        return jsonify({"ok": False, "error": "A user with that email already exists."}), 400

    requests_list = load_requests()

    # Check no pending request for same email
    existing = [r for r in requests_list
                if r.get("email") == email and r.get("status") == "pending"]
    if existing:
        return jsonify({"ok": False,
                        "error": "A pending request for that email already exists."}), 400

    import uuid
    new_request = {
        "id":           str(uuid.uuid4()),
        "name":         name,
        "email":        email,
        "role":         role,
        "note":         note,
        "submitted_by": username,
        "submitted_at": datetime.now().isoformat(),
        "status":       "pending",
        "decided_at":   None,
        "decided_by":   None,
    }
    requests_list.append(new_request)
    save_requests(requests_list)
    return jsonify({"ok": True, "request": new_request})


@app.route("/api/requests/approve", methods=["POST"])
@login_required
@admin_required
def approve_request():
    data       = request.get_json()
    request_id = data.get("id", "")

    requests_list = load_requests()
    req = next((r for r in requests_list if r["id"] == request_id), None)
    if not req:
        return jsonify({"ok": False, "error": "Request not found."}), 404
    if req["status"] != "pending":
        return jsonify({"ok": False, "error": "Request already decided."}), 400

    # Generate compliant temporary password
    special  = "!@#$%^&*"
    upper    = string.ascii_uppercase
    lower    = string.ascii_lowercase
    digits   = string.digits
    pwd_list = (
        [secrets.choice(upper)] +
        [secrets.choice(special)] +
        [secrets.choice(digits)] +
        [secrets.choice(upper + lower + digits + special) for _ in range(12)]
    )
    secrets.SystemRandom().shuffle(pwd_list)
    temp_password = "".join(pwd_list)

    # Create the user account
    users = load_users()
    email = req["email"]
    if email in users:
        return jsonify({"ok": False,
                        "error": "A user with that email already exists."}), 400

    users[email] = {
        "password":    _hash_password(temp_password),
        "role":        req["role"],
        "created_at":  datetime.now().isoformat(),
        "expires_at":  (datetime.now() + timedelta(days=14)).isoformat(),
        "must_change": True,
        "name":        req["name"],
        "active":      True,
    }
    save_users(users)

    # Update request status
    req["status"]     = "approved"
    req["decided_at"] = datetime.now().isoformat()
    req["decided_by"] = session.get("username", "")
    save_requests(requests_list)

    return jsonify({
        "ok":           True,
        "username":     email,
        "temp_password": temp_password,
        "name":         req["name"],
        "role":         req["role"],
    })


@app.route("/api/requests/deny", methods=["POST"])
@login_required
@admin_required
def deny_request():
    data       = request.get_json()
    request_id = data.get("id", "")
    reason     = data.get("reason", "").strip()

    requests_list = load_requests()
    req = next((r for r in requests_list if r["id"] == request_id), None)
    if not req:
        return jsonify({"ok": False, "error": "Request not found."}), 404
    if req["status"] != "pending":
        return jsonify({"ok": False, "error": "Request already decided."}), 400

    req["status"]     = "denied"
    req["reason"]     = reason
    req["decided_at"] = datetime.now().isoformat()
    req["decided_by"] = session.get("username", "")
    save_requests(requests_list)
    return jsonify({"ok": True})


@app.route("/admin/requests")
@login_required
@admin_required
def admin_requests():
    all_requests = load_requests()
    # Sort — pending first, then by date descending
    all_requests.sort(key=lambda r: (r["status"] != "pending",
                                      r["submitted_at"]), reverse=False)
    pending  = [r for r in all_requests if r["status"] == "pending"]
    decided  = [r for r in all_requests if r["status"] != "pending"]
    return render_template("admin_requests.html",
                           pending=pending, decided=decided,
                           username=session.get("username", ""),
                           role=session.get("role", ""))


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — CLIENT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@app.route("/clients")
@login_required
def client_list():
    role     = session.get("role", "")
    username = session.get("username", "")
    if role == "Representative":
        return redirect(url_for("rep_portal"))
    visible = clients_for_user(username, role)
    return render_template("client_list.html",
                           clients=visible,
                           role=role,
                           username=username)


@app.route("/clients/<client_id>")
@login_required
def client_profile(client_id):
    role     = session.get("role", "")
    username = session.get("username", "")
    if role == "Representative":
        return redirect(url_for("rep_portal"))

    client = get_client(client_id)
    if not client:
        return "Client not found.", 404

    # Access control — non-admins must be assigned
    if role != "Administrator" and username not in client.get("assigned_users", []):
        return "Access denied.", 403

    # Build summary from full history
    all_history = load_history()
    summary     = client_summary(client, all_history)

    # Load users for admin assignment panel
    users = load_users() if role == "Administrator" else {}
    safe_users = {u: d["role"] for u, d in users.items()
                  if d["role"] != "Representative"}

    return render_template("client_profile.html",
                           client=client,
                           summary=summary,
                           role=role,
                           username=username,
                           all_users=safe_users)


@app.route("/api/clients/create", methods=["POST"])
@login_required
@admin_required
def create_client():
    data = request.get_json()
    name = re.sub(r"[^a-zA-Z0-9 ]", "", data.get("name", "").strip())
    if not name:
        return jsonify({"ok": False, "error": "Client name is required (letters and numbers only)."}), 400

    clients = load_clients()
    # Check for duplicate name
    if any(c["name"].lower() == name.lower() for c in clients.values()):
        return jsonify({"ok": False, "error": "A client with that name already exists."}), 400

    import uuid
    client_id = str(uuid.uuid4())
    clients[client_id] = {
        "client_id":      client_id,
        "name":           name,
        "industry":       re.sub(r"[^a-zA-Z0-9 ]", "", data.get("industry", "").strip()),
        "contact_name":   re.sub(r"[^a-zA-Z0-9 ]", "", data.get("contact_name", "").strip()),
        "contact_email":  data.get("contact_email", "").strip(),
        "notes":          data.get("notes", "").strip()[:500],
        "created_at":     datetime.now().strftime("%Y-%m-%d"),
        "created_by":     session.get("username", ""),
        "assigned_users": [],
        "assessments":    [],
        "last_tested":    None,
    }
    save_clients(clients)
    return jsonify({"ok": True, "client_id": client_id, "name": name})


@app.route("/api/clients/update", methods=["POST"])
@login_required
@admin_required
def update_client():
    data      = request.get_json()
    client_id = data.get("client_id", "")
    clients   = load_clients()
    if client_id not in clients:
        return jsonify({"ok": False, "error": "Client not found."}), 404

    if "name" in data:
        name = re.sub(r"[^a-zA-Z0-9 ]", "", data["name"].strip())
        if not name:
            return jsonify({"ok": False, "error": "Name cannot be empty."}), 400
        clients[client_id]["name"] = name
    if "industry"      in data: clients[client_id]["industry"]      = re.sub(r"[^a-zA-Z0-9 ]", "", data["industry"].strip())
    if "contact_name"  in data: clients[client_id]["contact_name"]  = re.sub(r"[^a-zA-Z0-9 ]", "", data["contact_name"].strip())
    if "contact_email" in data: clients[client_id]["contact_email"] = data["contact_email"].strip()
    if "notes"         in data: clients[client_id]["notes"]         = data["notes"].strip()[:500]

    save_clients(clients)
    return jsonify({"ok": True})


@app.route("/api/clients/assign", methods=["POST"])
@login_required
@admin_required
def assign_client_users():
    data      = request.get_json()
    client_id = data.get("client_id", "")
    users_list = data.get("users", [])
    clients   = load_clients()
    if client_id not in clients:
        return jsonify({"ok": False, "error": "Client not found."}), 404
    clients[client_id]["assigned_users"] = users_list
    save_clients(clients)
    return jsonify({"ok": True})


@app.route("/api/clients/delete", methods=["POST"])
@login_required
@admin_required
def delete_client():
    data      = request.get_json()
    client_id = data.get("client_id", "")
    clients   = load_clients()
    if client_id not in clients:
        return jsonify({"ok": False, "error": "Client not found."}), 404
    del clients[client_id]
    save_clients(clients)
    return jsonify({"ok": True})


@app.route("/api/clients/search")
@login_required
def search_clients():
    """Search clients by name — alphanumeric only."""
    role     = session.get("role", "")
    username = session.get("username", "")
    query    = re.sub(r"[^a-zA-Z0-9 ]", "", request.args.get("q", "")).strip().lower()
    visible  = clients_for_user(username, role)
    if query:
        visible = [c for c in visible if query in c["name"].lower()]
    return jsonify({"ok": True, "clients": visible[:20]})


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — BREACH CHECK
# ══════════════════════════════════════════════════════════════════════════

@app.route("/breach")
@login_required
def breach_page():
    return render_template("breach_check.html",
                           role=session.get("role", ""),
                           username=session.get("username", ""))


@app.route("/api/breach/check", methods=["POST"])
@login_required
def breach_check_api():
    data   = request.get_json()
    target = data.get("target", "").strip().lower()
    target = re.sub(r"[^a-zA-Z0-9@._\-]", "", target)

    if not target:
        return jsonify({"ok": False, "error": "Enter a domain or email address."}), 400

    api_key = OATHNET_API_KEY
    if not api_key:
        return jsonify({"ok": False,
                        "error": "OATHNET_API_KEY not configured in service file."}), 500

    import requests as req_lib

    def call_oathnet(endpoint: str) -> tuple[list, int, str, int | None]:
        """Call an OathNet endpoint, return (items, total, error, lookups_left)."""
        try:
            resp = req_lib.get(
                f"https://oathnet.org/api/service/v2/{endpoint}/search",
                params={"q": target},
                headers={"x-api-key": api_key},
                timeout=15
            )
            body = resp.json()
            if not body.get("success"):
                if resp.status_code == 401:
                    return [], 0, "API key invalid or expired.", None
                if resp.status_code == 429:
                    return [], 0, "Rate limit hit — try again shortly.", None
                return [], 0, body.get("message", f"API error on {endpoint}"), None
            meta         = body.get("data", {}).get("meta", {})
            items        = body.get("data", {}).get("items", [])
            total        = meta.get("total", 0)
            lookups_left = meta.get("lookups_left", None)
            return items, total, "", lookups_left
        except Exception as ex:
            return [], 0, str(ex), None

    def process_breach(items: list) -> dict:
        seen = set()
        sources = set()
        out = []
        for i in items:
            if not isinstance(i, dict):
                continue
            sources.add(i.get("dbname", "unknown"))
            email        = i.get("email", "")
            username     = i.get("username", "")
            full_name    = i.get("full_name", "")
            email_domain = i.get("email_domain", "")
            key = f"{email}:{username}"
            if key not in seen:
                seen.add(key)
            out.append({
                "email":        email,
                "username":     username,
                "full_name":    full_name,
                "email_domain": email_domain,
                "dbname":       i.get("dbname", "unknown"),
            })
        return {"items": out, "unique": len(seen), "sources": list(sources)}

    def process_stealer(items: list) -> dict:
        seen = set()
        sources = set()
        out = []
        for i in items:
            if not isinstance(i, dict):
                continue
            sources.add(i.get("dbname", i.get("stealer", "unknown")))
            username = i.get("username", "")
            password = i.get("password", "")
            # pwned_at is the correct stealer field per OathNet docs
            pwned_at = i.get("pwned_at", "")
            if pwned_at and "T" in pwned_at:
                pwned_at = pwned_at.split("T")[0]
            key = f"{username}:{password}"
            if key not in seen:
                seen.add(key)
            out.append({
                "username": username,
                "password": password,
                "pwned_at": pwned_at,
                "dbname":   i.get("dbname", i.get("stealer", "unknown")),
            })
        return {"items": out, "unique": len(seen), "sources": list(sources)}

    # Call both endpoints
    breach_items, breach_total, breach_err, breach_lookups   = call_oathnet("breach")
    stealer_items, stealer_total, stealer_err, stealer_lookups = call_oathnet("stealer")

    # If both failed with auth error, return error
    if breach_err and stealer_err:
        return jsonify({"ok": False, "error": breach_err or stealer_err}), 500

    breach_data  = process_breach(breach_items)
    stealer_data = process_stealer(stealer_items)

    # Use whichever endpoint returned a lookups count
    lookups_left = breach_lookups if breach_lookups is not None else stealer_lookups

    return jsonify({
        "ok":          True,
        "target":      target,
        "lookups_left": lookups_left,
        "breach": {
            "total":   breach_total,
            "unique":  breach_data["unique"],
            "sources": breach_data["sources"],
            "items":   breach_data["items"],
            "error":   breach_err,
        },
        "stealer": {
            "total":   stealer_total,
            "unique":  stealer_data["unique"],
            "sources": stealer_data["sources"],
            "items":   stealer_data["items"],
            "error":   stealer_err,
        },
    })


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — FULL HISTORY
# ══════════════════════════════════════════════════════════════════════════

@app.route("/history")
@login_required
def full_history():
    username = session.get("username", "")
    role     = session.get("role", "")
    if role == "Administrator":
        # Admins can see all history, grouped by user
        history = load_history()
    else:
        history = load_history_for_user(username)
    return render_template("history.html",
                           history=history,
                           role=role,
                           username=username)


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — FINDINGS REPOSITORY
# ══════════════════════════════════════════════════════════════════════════

@app.route("/findings")
@login_required
def findings_repo():
    role     = session.get("role", "")
    username = session.get("username", "")
    if role == "Representative":
        return redirect(url_for("rep_portal"))
    findings = load_findings()
    # Sort by severity then title
    sorted_findings = sorted(
        findings.values(),
        key=lambda f: (SEVERITY_ORDER.get(f.get("severity", "Low"), 5),
                       f.get("title", "").lower())
    )
    return render_template("findings_repo.html",
                           findings=sorted_findings,
                           role=role,
                           username=username)


@app.route("/api/findings/create", methods=["POST"])
@login_required
def create_finding():
    role = session.get("role", "")
    if role not in ("Administrator", "Manager"):
        return jsonify({"ok": False, "error": "Access denied."}), 403

    data  = request.get_json()
    title = data.get("title", "").strip()
    if not title:
        return jsonify({"ok": False, "error": "Title is required."}), 400

    import uuid
    finding_id = str(uuid.uuid4())
    now        = datetime.now().strftime("%Y-%m-%d")

    findings = load_findings()
    findings[finding_id] = {
        "finding_id":         finding_id,
        "title":              title,
        "severity":           data.get("severity", "Medium"),
        "description":        data.get("description", "").strip(),
        "steps_to_reproduce": data.get("steps_to_reproduce", "").strip(),
        "affected_resource":  data.get("affected_resource", "").strip(),
        "references":         [r.strip() for r in
                               data.get("references", "").splitlines()
                               if r.strip()],
        "remediation":        data.get("remediation", "").strip(),
        "created_by":         session.get("username", ""),
        "created_at":         now,
        "updated_at":         now,
    }
    save_findings(findings)
    return jsonify({"ok": True, "finding_id": finding_id, "title": title})


@app.route("/api/findings/update", methods=["POST"])
@login_required
def update_finding():
    role = session.get("role", "")
    if role not in ("Administrator", "Manager"):
        return jsonify({"ok": False, "error": "Access denied."}), 403

    data       = request.get_json()
    finding_id = data.get("finding_id", "")
    findings   = load_findings()
    if finding_id not in findings:
        return jsonify({"ok": False, "error": "Finding not found."}), 404

    f = findings[finding_id]
    if "title"              in data: f["title"]              = data["title"].strip()
    if "severity"           in data: f["severity"]           = data["severity"]
    if "description"        in data: f["description"]        = data["description"].strip()
    if "steps_to_reproduce" in data: f["steps_to_reproduce"] = data["steps_to_reproduce"].strip()
    if "affected_resource"  in data: f["affected_resource"]  = data["affected_resource"].strip()
    if "remediation"        in data: f["remediation"]        = data["remediation"].strip()
    if "references"         in data:
        f["references"] = [r.strip() for r in
                           data["references"].splitlines() if r.strip()]
    f["updated_at"] = datetime.now().strftime("%Y-%m-%d")
    save_findings(findings)
    return jsonify({"ok": True})


@app.route("/api/findings/delete", methods=["POST"])
@login_required
def delete_finding():
    if session.get("role") != "Administrator":
        return jsonify({"ok": False, "error": "Admin only."}), 403
    data       = request.get_json()
    finding_id = data.get("finding_id", "")
    findings   = load_findings()
    if finding_id not in findings:
        return jsonify({"ok": False, "error": "Finding not found."}), 404
    del findings[finding_id]
    save_findings(findings)
    return jsonify({"ok": True})


@app.route("/api/findings/list")
@login_required
def list_findings():
    """Return all findings sorted by severity — used by assessment picker modal."""
    findings = load_findings()
    sorted_f = sorted(
        findings.values(),
        key=lambda f: (SEVERITY_ORDER.get(f.get("severity", "Low"), 5),
                       f.get("title", "").lower())
    )
    return jsonify({"ok": True, "findings": sorted_f})


# ══════════════════════════════════════════════════════════════════════════
# FLASK ROUTES — WIKI
# ══════════════════════════════════════════════════════════════════════════

def _render_markdown(text: str) -> str:
    """Convert markdown to HTML, installing markdown library if needed."""
    try:
        import markdown
    except ImportError:
        subprocess.run(["pip3", "install", "markdown", "--break-system-packages"],
                       capture_output=True)
        import markdown
    return markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "toc", "nl2br"]
    )


@app.route("/wiki")
@login_required
def wiki_home():
    wiki = get_or_init_wiki()
    # Group by phase
    phases = {}
    for slug, article in wiki.items():
        phase = article.get("phase", "Other")
        phases.setdefault(phase, []).append({**article, "slug": slug})
    # Sort phases in order
    phase_order = [
        "Phase 1 — Recon / OSINT",
        "Phase 2 — Port Scanning",
        "Phase 3 — Vulnerability Scanning",
        "Other"
    ]
    ordered = {p: phases[p] for p in phase_order if p in phases}
    for p in phases:
        if p not in ordered:
            ordered[p] = phases[p]
    return render_template("wiki_home.html",
                           phases=ordered,
                           role=session.get("role", ""),
                           username=session.get("username", ""))


@app.route("/wiki/<slug>")
@login_required
def wiki_article(slug):
    wiki = get_or_init_wiki()
    if slug not in wiki:
        return redirect(url_for("wiki_home"))
    article  = wiki[slug]
    rendered = _render_markdown(article["content"])
    return render_template("wiki_article.html",
                           slug=slug,
                           article=article,
                           rendered=rendered,
                           role=session.get("role", ""),
                           username=session.get("username", ""))


@app.route("/wiki/<slug>/edit", methods=["GET", "POST"])
@login_required
@admin_required
def wiki_edit(slug):
    wiki = get_or_init_wiki()
    if slug not in wiki:
        return redirect(url_for("wiki_home"))

    if request.method == "POST":
        data    = request.get_json()
        content = data.get("content", "").strip()
        wiki[slug]["content"]     = content
        wiki[slug]["last_edited"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        wiki[slug]["edited_by"]   = session.get("username", "")
        save_wiki(wiki)
        return jsonify({"ok": True})

    article = wiki[slug]
    return render_template("wiki_edit.html",
                           slug=slug,
                           article=article,
                           role=session.get("role", ""),
                           username=session.get("username", ""))


# ══════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

def startup():
    q = scan_state["log_queue"]

    def log_fn(msg, tag="info"):
        q.put({"ts": datetime.now().strftime("%H:%M:%S"), "msg": msg, "tag": tag})
        print(f"[{tag.upper()}] {msg}")

    # Bootstrap default admin if no users exist
    bootstrap_admin()

    log_fn("╔══════════════════════════════════════╗", "phase")
    log_fn("║   ARTEMIS WEB — Starting up...       ║", "phase")
    log_fn("╚══════════════════════════════════════╝", "phase")
    tool_paths = check_and_install_tools(log_fn)
    scan_state["tool_paths"] = tool_paths
    log_fn("🌐 Artemis Web ready on http://localhost:5000", "success")


if __name__ == "__main__":
    startup_thread = threading.Thread(target=startup, daemon=True)
    startup_thread.start()
    app.run(host="127.0.0.1", port=5000, debug=False, threaded=True)
