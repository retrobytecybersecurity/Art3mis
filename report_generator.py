#!/usr/bin/env python3
"""
report_generator.py — Artemis Report Generator
Produces a PDF and DOCX pentest summary report from scan results.
Dependencies: reportlab, python-docx
    pip install reportlab python-docx
"""

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Callable

# ── ReportLab imports ─────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# ── python-docx imports ───────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════
# STRING SANITIZATION
# ══════════════════════════════════════════════════════════════════════════

def _s(value) -> str:
    """
    Sanitize any value for safe use in PDF/DOCX output.
    - Converts to string
    - Strips null bytes and non-XML-compatible control characters
    - Keeps printable ASCII, tabs, newlines, and valid Unicode
    - Truncates extremely long strings to prevent layout issues
    """
    if value is None:
        return "—"
    text = str(value)
    # Remove null bytes
    text = text.replace("\x00", "")
    # Remove control characters except tab (\x09), newline (\x0a), carriage return (\x0d)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Remove invalid XML unicode characters (surrogates etc.)
    text = re.sub(
        r"[^\u0009\u000a\u000d\u0020-\uD7FF\uE000-\uFFFD]", "", text
    )
    # Truncate to avoid ReportLab overflow
    if len(text) > 2000:
        text = text[:2000] + "…"
    return text.strip() or "—"


def _sl(items: list, limit: int = 50) -> list:
    """Sanitize a list of strings, capping at limit items."""
    return [_s(i) for i in (items or [])[:limit]]

# ══════════════════════════════════════════════════════════════════════════
# COLOUR PALETTE (shared)
# ══════════════════════════════════════════════════════════════════════════
C_BG       = colors.HexColor("#0a0d14")
C_ACCENT   = colors.HexColor("#00d4ff")
C_ACCENT2  = colors.HexColor("#ff6b35")
C_SUCCESS  = colors.HexColor("#00ff88")
C_WARN     = colors.HexColor("#ffcc00")
C_ERROR    = colors.HexColor("#ff3355")
C_PANEL    = colors.HexColor("#0f1420")
C_BORDER   = colors.HexColor("#1e2a40")
C_TEXT     = colors.HexColor("#e8eaf0")
C_DIM      = colors.HexColor("#5a6a8a")
C_WHITE    = colors.white
C_BLACK    = colors.black

# ══════════════════════════════════════════════════════════════════════════
# PDF REPORT
# ══════════════════════════════════════════════════════════════════════════

def _pdf_styles():
    base = getSampleStyleSheet()

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    styles = {
        "cover_title": S("cover_title",
            fontName="Helvetica-Bold", fontSize=36,
            textColor=C_ACCENT, alignment=TA_CENTER, spaceAfter=6),
        "cover_sub": S("cover_sub",
            fontName="Helvetica", fontSize=14,
            textColor=C_TEXT, alignment=TA_CENTER, spaceAfter=4),
        "cover_meta": S("cover_meta",
            fontName="Helvetica", fontSize=11,
            textColor=C_DIM, alignment=TA_CENTER, spaceAfter=2),
        "section": S("section",
            fontName="Helvetica-Bold", fontSize=15,
            textColor=C_ACCENT, spaceBefore=18, spaceAfter=6),
        "subsection": S("subsection",
            fontName="Helvetica-Bold", fontSize=11,
            textColor=C_ACCENT2, spaceBefore=10, spaceAfter=4),
        "body": S("body",
            fontName="Helvetica", fontSize=9,
            textColor=C_TEXT, spaceAfter=4, leading=14),
        "body_mono": S("body_mono",
            fontName="Courier", fontSize=8,
            textColor=C_TEXT, spaceAfter=2, leading=12),
        "note": S("note",
            fontName="Helvetica-Oblique", fontSize=8,
            textColor=C_DIM, spaceAfter=2),
    }
    return styles


def _hr(color=C_BORDER, thickness=0.5):
    return HRFlowable(width="100%", thickness=thickness, color=color, spaceAfter=6)


def _table_style_main():
    return TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  C_PANEL),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  C_ACCENT),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, 0),  9),
        ("BACKGROUND",   (0, 1), (-1, -1), C_BG),
        ("TEXTCOLOR",    (0, 1), (-1, -1), C_TEXT),
        ("FONTNAME",     (0, 1), (-1, -1), "Courier"),
        ("FONTSIZE",     (0, 1), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [C_BG, C_PANEL]),
        ("GRID",         (0, 0), (-1, -1), 0.3, C_BORDER),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("WORDWRAP",     (0, 0), (-1, -1), True),
    ])


def build_pdf(results: dict, folder: Path, styles: dict) -> Path:
    client   = _s(results.get("client", "Unknown"))
    date     = _s(results.get("date",   datetime.now().strftime("%Y-%m-%d")))
    targets  = _sl(results.get("targets", []))
    ports    = {_s(k): _sl(v) for k, v in results.get("open_ports", {}).items()}
    vulns    = {_s(k): _sl(v, 30) for k, v in results.get("vulnerabilities", {}).items()}
    subs     = _sl(results.get("subdomains", []))
    headers  = {_s(k): _sl(v) for k, v in results.get("missing_headers", {}).items()}
    ffuf     = {_s(k): v for k, v in results.get("ffuf_findings", {}).items()}
    msf      = _sl(results.get("msf_findings", []))
    o365     = results.get("o365_findings", {})
    metagoofil_data = _sl(results.get("metagoofil", []))

    out_path = folder / f"Artemis_Report_{client}_{date}.pdf"
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch,  bottomMargin=0.75*inch,
    )

    S = styles
    story = []

    # ── COVER ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 1.2*inch))
    story.append(Paragraph("◈ ARTEMIS", S["cover_title"]))
    story.append(Paragraph("External Network Penetration Test Report", S["cover_sub"]))
    story.append(Spacer(1, 0.3*inch))
    story.append(_hr(C_ACCENT, 1.5))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Client:  {client}", S["cover_meta"]))
    story.append(Paragraph(f"Date:    {date}", S["cover_meta"]))
    story.append(Paragraph(f"Prepared by:  Artemis Automated Scan Suite", S["cover_meta"]))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(
        "⚠  CONFIDENTIAL — This document contains sensitive security information. "
        "Distribution is restricted to authorized personnel only.",
        ParagraphStyle("warn", fontName="Helvetica-Oblique", fontSize=9,
                       textColor=C_WARN, alignment=TA_CENTER)))
    story.append(PageBreak())

    # ── SECTION 1 — Vulnerability / Port Summary ──────────────────────
    story.append(Paragraph("1. Vulnerability & Port Summary", S["section"]))
    story.append(_hr())

    # Use Paragraph for vuln text so ReportLab wraps long content naturally
    wrap_style = ParagraphStyle("cell_wrap",
        fontName="Courier", fontSize=8, textColor=C_TEXT, leading=11)

    if targets:
        rows = [["Host / IP", "Vulnerabilities Found", "Open Ports"]]
        for t in targets:
            v_list = vulns.get(t, [])
            p_list = ports.get(t, [])
            v_text = "\n".join(_s(v) for v in v_list[:10]) if v_list else "None detected"
            if len(v_list) > 10:
                v_text += f"\n... (+{len(v_list)-10} more)"
            p_text = ", ".join(sorted(set(_s(p) for p in p_list), key=lambda x: int(x) if x.isdigit() else 0)) if p_list else "—"
            rows.append([
                Paragraph(_s(t), wrap_style),
                Paragraph(v_text.replace("\n", "<br/>"), wrap_style),
                Paragraph(p_text, wrap_style),
            ])
        tbl = Table(rows, colWidths=[1.8*inch, 3.8*inch, 1.8*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
    else:
        story.append(Paragraph("No targets scanned.", S["note"]))
    story.append(Spacer(1, 0.2*inch))

    # ── SECTION 2 — OSINT / Subdomains ───────────────────────────────
    story.append(Paragraph("2. OSINT — Subdomain & OSINT Enumeration", S["section"]))
    story.append(_hr())

    # 2a — Subdomains: summary count only
    story.append(Paragraph("2a. Subdomains Discovered", S["subsection"]))
    unique_subs = list(set(subs))
    if unique_subs:
        rows = [["Source", "Total Subdomains Discovered"],
                ["assetfinder + amass + bbot", str(len(unique_subs))]]
        tbl = Table(rows, colWidths=[3.5*inch, 3.5*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
    else:
        story.append(Paragraph("No subdomains discovered.", S["note"]))
    story.append(Spacer(1, 0.1*inch))

    story.append(Spacer(1, 0.2*inch))

    # ── SECTION 3 — Domain Security ──────────────────────────────────
    story.append(Paragraph("3. Domain Security", S["section"]))
    story.append(_hr())

    # 3a — spoofy
    story.append(Paragraph("3a. Email Spoofability (spoofy)", S["subsection"]))
    story.append(Paragraph(
        "See spoofy_*.txt in the 1_recon folder for full SPF/DMARC/DKIM analysis.",
        S["note"]))
    story.append(Spacer(1, 0.08*inch))

    # 3b — O365 / Azure tenant
    story.append(Paragraph("3b. O365 / Azure Tenant Enumeration (o365spray)", S["subsection"]))
    if o365:
        flags = []
        if o365.get("o365"):     flags.append("Microsoft 365 tenant detected")
        if o365.get("adfs"):     flags.append("ADFS / federated identity detected")
        if o365.get("exchange"): flags.append("Exchange Online detected")
        if not flags:            flags.append("No cloud services definitively identified")
        rows = [["Domain", "Findings"]]
        rows.append([o365.get("domain", "—"), "\n".join(flags)])
        tbl = Table(rows, colWidths=[2.5*inch, 4.5*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
    else:
        story.append(Paragraph("o365spray not run or no findings.", S["note"]))
    story.append(Spacer(1, 0.2*inch))

    # ── SECTION 4 — Missing HTTP Security Headers ─────────────────────
    story.append(Paragraph("4. Missing HTTP Security Headers", S["section"]))
    story.append(_hr())
    if headers:
        rows = [["Host", "Missing Header Count"]]
        for host, h_list in headers.items():
            rows.append([host, str(len(h_list))])
        tbl = Table(rows, colWidths=[4.5*inch, 2.5*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
        story.append(Paragraph(
            "See shcheck_*.txt in 3_vuln/ for full header details.", S["note"]))
    else:
        story.append(Paragraph("No missing security headers detected.", S["note"]))
    story.append(Spacer(1, 0.2*inch))

    # ── SECTION 5 — Web Content Discovery (FFUF) ─────────────────────
    story.append(Paragraph("5. Web Content Discovery (FFUF — HTTP 200)", S["section"]))
    story.append(_hr())
    if ffuf:
        rows = [["Host", "HTTP 200 Paths Found"]]
        for host, findings in ffuf.items():
            count = len([f for f in findings if f.get("status") == 200])
            rows.append([host, str(count)])
        tbl = Table(rows, colWidths=[4.5*inch, 2.5*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
        story.append(Paragraph(
            "See ffuf_*.json in 3_vuln/ for full URL listing.", S["note"]))
    else:
        story.append(Paragraph("No HTTP 200 responses discovered by FFUF.", S["note"]))
    story.append(Spacer(1, 0.2*inch))

    # ── SECTION 6 — Metasploit Findings ──────────────────────────────
    story.append(Paragraph("6. Metasploit Auxiliary Findings", S["section"]))
    story.append(_hr())
    if msf:
        rows = [["Finding"]] + [[m] for m in msf]
        tbl = Table(rows, colWidths=[7.0*inch], repeatRows=1)
        tbl.setStyle(_table_style_main())
        story.append(tbl)
    else:
        story.append(Paragraph("No notable findings from Metasploit auxiliary modules.", S["note"]))

    # ── Footer ────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.4*inch))
    story.append(_hr(C_ACCENT, 0.8))
    story.append(Paragraph(
        f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        "Artemis Automated Pentest Suite  |  CONFIDENTIAL",
        ParagraphStyle("footer", fontName="Helvetica", fontSize=8,
                       textColor=C_DIM, alignment=TA_CENTER)))

    doc.build(story)
    return out_path


# ══════════════════════════════════════════════════════════════════════════
# DOCX REPORT
# ══════════════════════════════════════════════════════════════════════════

def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _heading_para(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = _rgb("#00d4ff")
    else:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = _rgb("#ff6b35")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _add_table(doc: Document, rows: list[list[str]],
               col_widths: list[float] | None = None):
    """Add a styled table. rows[0] = header row."""
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(cell_text))
            run.font.name = "Courier New"
            run.font.size = Pt(8)

            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = _rgb("#00d4ff")
                _set_cell_bg(cell, "#0f1420")
            else:
                run.font.color.rgb = _rgb("#e8eaf0")
                bg = "#0a0d14" if r_idx % 2 == 1 else "#0f1420"
                _set_cell_bg(cell, bg)

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for c_idx, cell in enumerate(row.cells):
                if c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])

    return tbl


def build_docx(results: dict, folder: Path) -> Path:
    client  = _s(results.get("client", "Unknown"))
    date    = _s(results.get("date",   datetime.now().strftime("%Y-%m-%d")))
    targets = _sl(results.get("targets", []))
    ports   = {_s(k): _sl(v) for k, v in results.get("open_ports", {}).items()}
    vulns   = {_s(k): _sl(v, 30) for k, v in results.get("vulnerabilities", {}).items()}
    subs    = _sl(results.get("subdomains", []))
    headers = {_s(k): _sl(v) for k, v in results.get("missing_headers", {}).items()}
    ffuf    = {_s(k): v for k, v in results.get("ffuf_findings", {}).items()}
    msf     = _sl(results.get("msf_findings", []))
    o365    = results.get("o365_findings", {})
    metagoofil_data = _sl(results.get("metagoofil", []))

    doc = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)
    section.top_margin    = Inches(0.85)
    section.bottom_margin = Inches(0.85)

    # ── Cover ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◈ ARTEMIS")
    run.font.size = Pt(32); run.font.bold = True
    run.font.color.rgb = _rgb("#00d4ff"); run.font.name = "Courier New"

    p2 = doc.add_paragraph("External Network Penetration Test Report")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14)
    p2.runs[0].font.color.rgb = _rgb("#e8eaf0")

    doc.add_paragraph()
    for label, value in [("Client", client), ("Date", date),
                          ("Prepared by", "Artemis Automated Scan Suite")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  {value}")
        r.font.size = Pt(11); r.font.color.rgb = _rgb("#5a6a8a")
        r.font.name = "Courier New"
    doc.add_page_break()

    # ── Section 1 — Vuln / Port Summary ──────────────────────────────
    _heading_para(doc, "1. Vulnerability & Port Summary")
    if targets:
        rows = [["Host / IP", "Vulnerabilities Found", "Open Ports"]]
        for t in targets:
            v_list = vulns.get(t, [])
            p_list = ports.get(t, [])
            v_text = "; ".join(_s(v) for v in v_list[:8]) if v_list else "None detected"
            if len(v_list) > 8: v_text += f" (+{len(v_list)-8} more)"
            p_text = ", ".join(sorted(set(_s(p) for p in p_list), key=lambda x: int(x) if x.isdigit() else 0)) if p_list else "—"
            t_display = (_s(t)[:40] + "...") if len(_s(t)) > 40 else _s(t)
            rows.append([t_display, v_text, p_text])
        _add_table(doc, rows, col_widths=[1.8, 3.8, 1.8])
    else:
        doc.add_paragraph("No targets scanned.")

    # ── Section 2 — OSINT ────────────────────────────────────────────
    _heading_para(doc, "2. OSINT — Subdomain & OSINT Enumeration")

    _heading_para(doc, "2a. Subdomains Discovered", level=2)
    unique_subs = list(set(subs))
    if unique_subs:
        rows = [["Source", "Total Subdomains Discovered"],
                ["assetfinder + amass + bbot", str(len(unique_subs))]]
        _add_table(doc, rows, col_widths=[3.5, 3.5])
        doc.add_paragraph("See assetfinder_*.txt, amass_*.txt and bbot_*.txt in 1_recon/ for full listing.")
    else:
        doc.add_paragraph("No subdomains discovered.")

    _heading_para(doc, "2b. metagoofil — Exposed File Metadata", level=2)
    if metagoofil_data:
        rows = [["File Types Searched", "Total Exposed Files Found"],
                ["pdf, doc, docx, xls, xlsx, ppt, pptx (root domain + subdomains)", str(len(metagoofil_data))]]
        _add_table(doc, rows, col_widths=[4.5, 2.5])
        doc.add_paragraph("See metagoofil_*.txt in 1_recon/ for full file listing.")
    else:
        doc.add_paragraph("metagoofil not run or no exposed files found.")

    # ── Section 3 — Domain Security ──────────────────────────────────
    _heading_para(doc, "3. Domain Security")

    _heading_para(doc, "3a. Email Spoofability (spoofy)", level=2)
    doc.add_paragraph("See spoofy_*.txt in the 1_recon folder for full SPF/DMARC/DKIM analysis.")

    _heading_para(doc, "3b. O365 / Azure Tenant Enumeration (o365spray)", level=2)
    if o365:
        flags = []
        if o365.get("o365"):     flags.append("Microsoft 365 tenant detected")
        if o365.get("adfs"):     flags.append("ADFS / federated identity detected")
        if o365.get("exchange"): flags.append("Exchange Online detected")
        if not flags:            flags.append("No cloud services definitively identified")
        rows = [["Domain", "Findings"]]
        rows.append([o365.get("domain", "—"), "; ".join(flags)])
        _add_table(doc, rows, col_widths=[2.5, 4.5])
    else:
        doc.add_paragraph("o365spray not run or no findings.")

    # ── Section 4 — Missing Headers ──────────────────────────────────
    _heading_para(doc, "4. Missing HTTP Security Headers")
    if headers:
        rows = [["Host", "Missing Header Count"]]
        for host, h_list in headers.items():
            rows.append([_s(host), str(len(h_list))])
        _add_table(doc, rows, col_widths=[4.5, 2.5])
        doc.add_paragraph("See shcheck_*.txt in 3_vuln/ for full header details.")
    else:
        doc.add_paragraph("No missing security headers detected.")

    # ── Section 5 — FFUF ─────────────────────────────────────────────
    _heading_para(doc, "5. Web Content Discovery (FFUF — HTTP 200)")
    if ffuf:
        rows = [["Host", "HTTP 200 Paths Found"]]
        for host, findings in ffuf.items():
            count = len([f for f in findings if isinstance(f, dict) and f.get("status") == 200])
            rows.append([_s(host), str(count)])
        _add_table(doc, rows, col_widths=[4.5, 2.5])
        doc.add_paragraph("See ffuf_*.txt in 3_vuln/ for full URL listing.")
    else:
        doc.add_paragraph("No HTTP 200 responses discovered by FFUF.")

    # ── Section 6 — Metasploit ────────────────────────────────────────
    _heading_para(doc, "6. Metasploit Auxiliary Findings")
    if msf:
        rows = [["Finding"]] + [[_s(m)] for m in msf]
        _add_table(doc, rows, col_widths=[7.0])
    else:
        doc.add_paragraph("No notable findings from Metasploit auxiliary modules.")

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph(
        f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
        "Artemis Automated Pentest Suite  |  CONFIDENTIAL"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = _rgb("#5a6a8a")

    out_path = folder / f"Artemis_Report_{client}_{date}.docx"
    doc.save(str(out_path))
    return out_path


# ══════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

def generate_reports(results: dict, folder: Path,
                     log_fn: Callable = print):
    styles = _pdf_styles()
    log_fn("  Generating PDF report...", "info")
    pdf_path = build_pdf(results, folder, styles)
    log_fn(f"  ✓ PDF: {pdf_path.name}", "success")

    log_fn("  Generating Word report...", "info")
    docx_path = build_docx(results, folder)
    log_fn(f"  ✓ DOCX: {docx_path.name}", "success")

    return pdf_path, docx_path
